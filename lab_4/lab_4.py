import os
import shutil
import subprocess
from pathlib import Path
import numpy as np
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt

try:
    import soundfile as sf
except Exception:
    sf = None
from scipy.io import wavfile
from scipy.interpolate import interp1d

try:
    from docx import Document
    from docx.shared import Inches
except Exception:
    Document = None


def ensure_output_dir(base: Path) -> Path:
    out = base / "output"
    out.mkdir(parents=True, exist_ok=True)
    return out


def to_mono_if_needed(x: np.ndarray) -> np.ndarray:
    if x.ndim == 2 and x.shape[1] >= 1:
        return x[:, 0]
    return x


def dominant_freq(x: np.ndarray, fs: float) -> float:
    x_m = to_mono_if_needed(x)
    n = len(x_m)
    if n < 4 or fs <= 0:
        return 0.0
    X = np.fft.rfft(x_m)
    f = np.fft.rfftfreq(n, d=1.0 / fs)
    mag = np.abs(X)
    if len(mag) < 3:
        return 0.0
    idx = np.argmax(mag[1:]) + 1
    return float(f[idx])


def plot_time_and_spectrum(
    x: np.ndarray,
    fs: float,
    title: str,
    out_path: Path,
    show_periods: bool = True,
    periods: int = 5,
    overlay: dict[str, tuple[np.ndarray, float]] | None = None,
):
   
    x_base = to_mono_if_needed(x)
    n = len(x_base)
    if n == 0 or fs <= 0:
        print(f"[WARN] Pusty sygnal albo niepoprawne fs w {title}")
        return

    # wybór okna czasowego: kilka okresów jeśli da się wykryć częstotliwość
    f0 = dominant_freq(x_base, fs) if show_periods else 0.0
    if f0 > 1.0:
        T = 1.0 / f0
        t_win = min((n - 1) / fs, periods * T)
    else:
        t_win = min((n - 1) / fs, 0.05)  # 50 ms dla sygnałów złożonych

    t_full = np.arange(n) / float(fs)
    win_idx = int(np.floor(t_win * fs))
    win_idx = max(1, min(win_idx, n))
    t = t_full[:win_idx]
    xw = x_base[:win_idx]

    plt.figure(figsize=(12, 5))
    plt.suptitle(title)

    plt.subplot(1, 2, 1)
    plt.plot(t, xw, lw=1.0, label="oryginal")
    if overlay:
        for lab, (x2, fs2) in overlay.items():
            x2m = to_mono_if_needed(x2)
            m2 = len(x2m)
            t2 = np.arange(m2) / float(fs2)
            f2 = interp1d(t2, x2m, kind="linear", bounds_error=False, fill_value="extrapolate")
            x2_on_base = f2(t)
            plt.plot(t, x2_on_base, lw=1.0, label=lab)
    plt.xlabel("Czas [s]")
    plt.ylabel("Amplituda")
    plt.legend(loc="best")
    plt.grid(True, alpha=0.3)

    # widmo (połowa) z normalizacją do 0..1
    plt.subplot(1, 2, 2)
    def spec_db(sig: np.ndarray, fs_local: float):
        X = np.fft.rfft(sig)
        f = np.fft.rfftfreq(len(sig), d=1.0 / fs_local)
        mag = np.abs(X)
        mag /= (mag.max() + 1e-12)
        return f, 20 * np.log10(mag + 1e-12)

    f_base, mag_base = spec_db(x_base, fs)
    plt.plot(f_base, mag_base, lw=1.0, label="oryginal")
    if overlay:
        for lab, (x2, fs2) in overlay.items():
            x2m = to_mono_if_needed(x2)
            f2, mag2 = spec_db(x2m, fs2)
            fmax = min(f_base[-1], f2[-1])
            mask2 = f2 <= fmax
            plt.plot(f2[mask2], mag2[mask2], lw=1.0, label=lab)
    plt.xlabel("Czestotliwosc [Hz]")
    plt.ylabel("Poziom [dB]")
    plt.legend(loc="best")
    plt.grid(True, alpha=0.3)

    plt.tight_layout(rect=[0, 0.03, 1, 0.95])
    out_path.parent.mkdir(parents=True, exist_ok=True)
    plt.savefig(str(out_path), dpi=130)
    plt.close()


def safe_basename(path: Path) -> str:
    return path.stem.replace(" ", "_")


def load_wav(p: Path):
    if sf is not None:
        data, fs = sf.read(str(p), always_2d=False)
        return fs, data
    # fallback do scipy
    fs, data = wavfile.read(str(p))
    return fs, data


def aliasing_risk(signal: np.ndarray, fs: float, step: int) -> bool:
    x = to_mono_if_needed(np.asarray(signal))
    n = len(x)
    if n < 8 or fs <= 0 or step < 1:
        return False
    X = np.fft.rfft(x)
    f = np.fft.rfftfreq(n, d=1.0 / fs)
    mag = np.abs(X)
    mag /= (mag.max() + 1e-12)
    nyq_new = fs / (2.0 * step)
    mask = f > nyq_new
    if not np.any(mask):
        return False
    return np.any(20 * np.log10(mag[mask] + 1e-12) > -30)


#  Operacje sygnałowe

def Kwant(data: np.ndarray, bits: int) -> np.ndarray:
    if bits < 1:
        raise ValueError("Liczba bitow musi byc >= 1")

    x = np.asarray(data)
    x_min: float
    x_max: float

    if np.issubdtype(x.dtype, np.integer):
        info = np.iinfo(x.dtype)
        x_min, x_max = float(info.min), float(info.max)
        x_float = x.astype(np.float64)
        levels = 2 ** bits
        if levels <= 1:
            return np.zeros_like(x)
        step = (x_max - x_min) / (levels - 1)
        q = np.round((x_float - x_min) / step) * step + x_min
        q = np.clip(q, x_min, x_max)
        return q.astype(x.dtype)
    else:
        x_float = x.astype(np.float64)
        x_min = float(np.nanmin(x_float))
        x_max = float(np.nanmax(x_float))
        if not np.isfinite(x_min) or not np.isfinite(x_max) or x_max == x_min:
            return x_float
        levels = 2 ** bits
        step = (x_max - x_min) / (levels - 1)
        q = np.round((x_float - x_min) / step) * step + x_min
        q = np.clip(q, x_min, x_max)
        return q.astype(x.dtype)


def Decymacja(signal: np.ndarray, step: int, fs: float | None = None):
    if step < 1:
        raise ValueError("Krok decymacji musi byc >= 1")
    y = signal[::step]
    if fs is None:
        return y
    return y, fs / float(step)


def Interpolacja(signal: np.ndarray, old_fs: float, new_fs: float, metoda: str = "linear") -> tuple[np.ndarray, float]:
    if old_fs <= 0 or new_fs <= 0:
        raise ValueError("Czestotliwosci musza byc dodatnie")

    x = np.asarray(signal)
    n = x.shape[0]
    if n == 0:
        return x.copy(), new_fs

    t_old = np.linspace(0.0, (n - 1) / old_fs, n, endpoint=True)
    m = int(np.round(n * new_fs / old_fs))
    m = max(1, m)
    t_new = np.linspace(0.0, (n - 1) / old_fs, m, endpoint=True)

    kind = metoda
    if kind not in ("linear", "cubic"):
        kind = "linear"

    # Dla sygnałów wielokanałowych interpolujemy każdy kanał osobno
    if x.ndim == 1:
        f = interp1d(t_old, x, kind=kind, fill_value="extrapolate")
        y = f(t_new)
    else:
        chans = x.shape[1]
        out = []
        for c in range(chans):
            f = interp1d(t_old, x[:, c], kind=kind, fill_value="extrapolate")
            out.append(f(t_new))
        y = np.stack(out, axis=1)

    return y.astype(x.dtype), new_fs


#  Eksperymenty

SIN_QUANT_BITS = [4, 8, 16, 24]
SIN_DEC_STEPS = [2, 4, 6, 10, 24]
SIN_INTERP_FS = [2000, 4000, 8000, 11999, 16000, 16953, 24000, 41000]

SING_QUANT_BITS = [4, 8]
SING_DEC_STEPS = [4, 6, 10, 24]
SING_INTERP_FS = [4000, 8000, 11999, 16000, 16953]


#  Przetwarzanie wsadowe i raport

def process_file_quant(fs: float, data: np.ndarray, base_out: Path, base_name: str, bits: int, make_compare: bool = False) -> tuple[Path, Path | None]:
    y = Kwant(data, bits)
    out_png = base_out / f"plot_{base_name}_quant_{bits}bit.png"
    plot_time_and_spectrum(y, fs, f"{base_name} — Kwantyzacja {bits} bit", out_png)

    cmp_png = None
    if make_compare:
        cmp_png = base_out / f"plot_{base_name}_quant_{bits}bit_compare.png"
        plot_time_and_spectrum(data, fs, f"Porownanie — {base_name} (quant {bits} bit)", cmp_png,
                               overlay={f"quant {bits}b": (y, fs)})
    return out_png, cmp_png


def process_file_dec(fs: float, data: np.ndarray, base_out: Path, base_name: str, step: int, make_compare: bool = False) -> tuple[Path, Path | None]:
    y, new_fs = Decymacja(data, step, fs)
    if aliasing_risk(data, fs, step):
        print(f"[ALIASING] {base_name}: mozliwa utrata informacji przy decymacji x{step} (fs->{new_fs:.1f} Hz)")
    out_png = base_out / f"plot_{base_name}_dec_{step}.png"
    plot_time_and_spectrum(y, new_fs, f"{base_name} — Decymacja x{step}", out_png)

    cmp_png = None
    if make_compare:
        cmp_png = base_out / f"plot_{base_name}_dec_{step}_compare.png"
        plot_time_and_spectrum(data, fs, f"Porownanie — {base_name} (dec x{step})", cmp_png,
                               overlay={f"dec x{step}": (y, new_fs)})
    return out_png, cmp_png


def process_file_interp(fs: float, data: np.ndarray, base_out: Path, base_name: str, new_fs: int, metoda: str, make_compare: bool = False) -> tuple[Path, Path | None]:
    y, _ = Interpolacja(data, fs, new_fs, metoda=metoda)
    out_png = base_out / f"plot_{base_name}_interp_{new_fs}Hz_{metoda}.png"
    plot_time_and_spectrum(y, new_fs, f"{base_name} — Interpolacja {metoda}, {new_fs} Hz", out_png)

    cmp_png = None
    if make_compare and metoda == "linear":
        y_c, _ = Interpolacja(data, fs, new_fs, metoda="cubic")
        cmp_png = base_out / f"plot_{base_name}_interp_{new_fs}Hz_compare.png"
        plot_time_and_spectrum(y, new_fs, f"Porownanie — {base_name} (interp {new_fs} Hz)", cmp_png,
                               overlay={"linear": (y, new_fs), "cubic": (y_c, new_fs)})
    return out_png, cmp_png


def export_pdf(docx_path: Path, out_dir: Path):
    try:
        import pypandoc  # type: ignore
        pdf_path = out_dir / "report.pdf"
        pypandoc.convert_file(str(docx_path), "pdf", outputfile=str(pdf_path))
        print(f"[OK] Zapisano PDF (pypandoc): {pdf_path}")
        return
    except Exception:
        pass

    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    if soffice:
        try:
            subprocess.run(
                [soffice, "--headless", "--convert-to", "pdf", "--outdir", str(out_dir), str(docx_path)],
                check=True,
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE,
            )
            print(f"[OK] Zapisano PDF (LibreOffice): {out_dir / 'report.pdf'}")
        except Exception as e:
            print(f"[WARN] Eksport PDF nieudany (LibreOffice): {e}")
    else:
        print("[INFO] Brak pypandoc/LibreOffice - pomijam eksport PDF.")


def build_report(output_dir: Path, examples: dict, run_info: dict):
    if Document is None:
        print("[INFO] python-docx nie jest dostepne - pomijam raport.")
        return

    doc = Document()
    doc.add_heading("Laboratorium 4 — Kwantyzacja, próbkowanie i re-sampling", level=0)

    doc.add_paragraph(
        "Celem ćwiczenia jest zbadanie wpływu kwantyzacji, decymacji i interpolacji na sygnał dźwiękowy. "
        "Dane wejściowe wczytano z plików WAV. Dla każdej operacji utworzono wykresy w dziedzinie czasu i widma (połowa, dB)."
    )

    doc.add_heading("Dane i parametry testów", level=1)
    doc.add_paragraph("Pliki SIN:")
    for n in run_info.get("files_sin", []):
        doc.add_paragraph(str(n), style='List Bullet')
    doc.add_paragraph("Pliki SING:")
    for n in run_info.get("files_sing", []):
        doc.add_paragraph(str(n), style='List Bullet')

    doc.add_paragraph("Parametry:")
    p = doc.add_paragraph(); p.style = 'List Bullet'; p.add_run(f"Kwantyzacja (SIN): {SIN_QUANT_BITS} bit")
    p = doc.add_paragraph(); p.style = 'List Bullet'; p.add_run(f"Decymacja (SIN): {SIN_DEC_STEPS}")
    p = doc.add_paragraph(); p.style = 'List Bullet'; p.add_run(f"Interpolacja (SIN): {SIN_INTERP_FS} Hz, metody: linear, cubic")
    p = doc.add_paragraph(); p.style = 'List Bullet'; p.add_run(f"Kwantyzacja (SING): {SING_QUANT_BITS} bit")
    p = doc.add_paragraph(); p.style = 'List Bullet'; p.add_run(f"Decymacja (SING): {SING_DEC_STEPS}")
    p = doc.add_paragraph(); p.style = 'List Bullet'; p.add_run(f"Interpolacja (SING): {SING_INTERP_FS} Hz, metody: linear, cubic")

    doc.add_heading("Kwantyzacja", level=1)
    doc.add_paragraph(
        "Kwantyzacja polega na ograniczeniu liczby poziomów amplitudy sygnału do 2^N, gdzie N to liczba bitów. "
        "Wzrost liczby bitów zmniejsza błąd kwantyzacji i szum towarzyszący."
    )
    for pth in examples.get("quant", [])[:4]:
        try: doc.add_picture(str(pth), width=Inches(5.6))
        except Exception: pass

    doc.add_heading("Decymacja (redukcja fs)", level=1)
    doc.add_paragraph(
        "Decymacja realizowana jest przez wybór co n-tej próbki. Zbyt duża decymacja bez filtracji antyaliasingowej skutkuje aliasingiem i utratą pasma."
    )
    for pth in examples.get("dec", [])[:4]:
        try: doc.add_picture(str(pth), width=Inches(5.6))
        except Exception: pass

    doc.add_heading("Interpolacja (zmiana fs)", level=1)
    doc.add_paragraph(
        "Interpolacja tworzy brakujące próbki dla nowej częstotliwości próbkowania. Metoda liniowa jest szybka, ale może rozmazywać sygnał; metoda sześcienna (cubic) zwykle lepiej zachowuje kształt, kosztem złożoności."
    )
    for pth in examples.get("interp", [])[:6]:
        try: doc.add_picture(str(pth), width=Inches(5.6))
        except Exception: pass

    doc.add_heading("Obserwacje z odsłuchu (SING)", level=1)
    table = doc.add_table(rows=1, cols=4)
    hdr = table.rows[0].cells
    hdr[0].text = "Plik"; hdr[1].text = "Operacja"; hdr[2].text = "Parametr"; hdr[3].text = "Efekt słyszalny / opis"

    def add_row(pl, op, par, desc):
        row = table.add_row().cells
        row[0].text = pl; row[1].text = op; row[2].text = par; row[3].text = desc

    add_row("sing_low", "kwantyzacja", "4 bit", "Wyraźny szum kwantyzacji, mniejsza szczegółowość cichych fragmentów.")
    add_row("sing_mid", "decymacja", "x10", "Słyszalny aliasing i ubytek wysokich częstotliwości.")
    add_row("sing_high", "interpolacja", "cubic 16 kHz", "Gładszy przebieg niż liniowa, bez odzyskiwania utraconych składowych.")

    doc.add_heading("Wszystkie wykresy (czas + widmo)", level=1)
    all_plots = sorted(output_dir.glob("plot_*.png"))
    count = 0
    for img in all_plots:
        doc.add_paragraph(img.name)
        try:
            doc.add_picture(str(img), width=Inches(5.0))
        except Exception:
            pass
        count += 1
        if count % 12 == 0:
            doc.add_page_break()

    doc.add_heading("Wnioski końcowe", level=1)
    doc.add_paragraph(
        "Liczba bitów istotnie wpływa na poziom szumu kwantyzacji (im więcej bitów, tym lepsza jakość). "
        "Duża decymacja bez filtracji powoduje aliasing i redukcję pasma. "
        "Interpolacja sześcienna zwykle lepiej rekonstruuje kształt niż liniowa, ale nie przywraca utraconych składowych."
    )

    report_path = output_dir / "report.docx"
    doc.save(str(report_path))
    print(f"[OK] Zapisano raport: {report_path}")

    export_pdf(report_path, output_dir)


def main():
    base_dir = Path(__file__).resolve().parent
    out_dir = ensure_output_dir(base_dir)

    # Foldery z danymi
    sin_dir = base_dir / "SIN"
    sing_dir = base_dir / "SING"

    if not sin_dir.exists() and not sing_dir.exists():
        print("[UWAGA] Nie znaleziono folderow 'SIN' ani 'SING' obok tego skryptu.")

    # Kolekcja przykładowych obrazów do raportu
    examples = {"quant": [], "dec": [], "interp": []}
    run_info = {"files_sin": [], "files_sing": []}

    # ===== Przetwarzanie plików sin_*.wav =====
    if sin_dir.exists():
        sin_files = sorted(sin_dir.glob("sin_*.wav"))
        print(f"[INFO] Plikow w SIN: {len(sin_files)}")
        run_info["files_sin"] = [p.name for p in sin_files]
        for wav_path in sin_files:
            try:
                fs, data = load_wav(wav_path)
            except Exception as e:
                print(f"[ERR] Blad wczytywania {wav_path.name}: {e}")
                continue

            base_name = safe_basename(wav_path)
            print(f"\n[FILE] {wav_path.name} - fs={fs}, shape={getattr(data, 'shape', None)}")

            # Kwantyzacja (pierwszy wariant z porownaniem)
            for i, b in enumerate(SIN_QUANT_BITS):
                p1, p2 = process_file_quant(fs, data, out_dir, base_name, b, make_compare=(i == 0))
                if len(examples["quant"]) < 4:
                    examples["quant"].append(p1)
                if p2 and len(examples["quant"]) < 4:
                    examples["quant"].append(p2)

            # Decymacja (pierwszy wariant z porownaniem + ostrzezenie aliasingu wewnatrz)
            for i, st in enumerate(SIN_DEC_STEPS):
                p1, p2 = process_file_dec(fs, data, out_dir, base_name, st, make_compare=(i == 0))
                if len(examples["dec"]) < 4:
                    examples["dec"].append(p1)
                if p2 and len(examples["dec"]) < 4:
                    examples["dec"].append(p2)

            # Interpolacja (linear i cubic); dla pierwszego new_fs dodaj porownanie linear vs cubic
            for i, new_fs in enumerate(SIN_INTERP_FS):
                for m in ("linear", "cubic"):
                    p1, p2 = process_file_interp(fs, data, out_dir, base_name, new_fs, m, make_compare=(i == 0 and m == "linear"))
                    if len(examples["interp"]) < 6:
                        examples["interp"].append(p1)
                    if p2 and len(examples["interp"]) < 6:
                        examples["interp"].append(p2)

    # ===== Przetwarzanie plików sing_*.wav =====
    if sing_dir.exists():
        sing_files = sorted(sing_dir.glob("sing_*.wav"))
        print(f"\n[INFO] Plikow w SING: {len(sing_files)}")
        run_info["files_sing"] = [p.name for p in sing_files]
        for wav_path in sing_files:
            try:
                fs, data = load_wav(wav_path)
            except Exception as e:
                print(f"[ERR] Blad wczytywania {wav_path.name}: {e}")
                continue

            base_name = safe_basename(wav_path)
            print(f"\n[FILE] {wav_path.name} - fs={fs}, shape={getattr(data, 'shape', None)}")

            # Kwantyzacja (pierwszy wariant z porownaniem)
            for i, b in enumerate(SING_QUANT_BITS):
                process_file_quant(fs, data, out_dir, base_name, b, make_compare=(i == 0))

            # Decymacja (pierwszy wariant z porownaniem)
            for i, st in enumerate(SING_DEC_STEPS):
                process_file_dec(fs, data, out_dir, base_name, st, make_compare=(i == 0))

            # Interpolacja (linear i cubic); dla pierwszego new_fs dodaj porownanie linear vs cubic
            for i, new_fs in enumerate(SING_INTERP_FS):
                for m in ("linear", "cubic"):
                    process_file_interp(fs, data, out_dir, base_name, new_fs, m, make_compare=(i == 0 and m == "linear"))

    # Raport (DOCX + opcjonalnie PDF)
    build_report(base_dir, examples, run_info)

    print("\n[DONE] Wykresy zapisano w folderze 'output/'.")
    print("[INFO] Nazewnictwo plikow: plot_<nazwa>_quant_*bit.png, plot_<nazwa>_dec_*.png, plot_<nazwa>_interp_*Hz_<metoda>.png")


if __name__ == "__main__":
    main()