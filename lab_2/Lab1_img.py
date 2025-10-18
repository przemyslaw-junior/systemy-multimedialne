import os
import numpy as np
import matplotlib.pyplot as plt
import cv2
import pandas as pd
from docx import Document
from docx.shared import Inches
from io import BytesIO

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
IMG_DIR = os.path.join(BASE_DIR, "images")
OUT_DIR = os.path.join(BASE_DIR, "output")

os.makedirs(IMG_DIR, exist_ok=True)
os.makedirs(OUT_DIR, exist_ok=True)

def imgToUInt8(img):
    if np.issubdtype(img.dtype, np.unsignedinteger):
        return img
    else:
        print("Konwersja z float do uint8")
        img_uint8 = (img * 255).astype('uint8')
        return img_uint8


def imgToFloat(img):
    if np.issubdtype(img.dtype, np.floating):
        return img
    else:
        print("Konwersja z uint8 do float")
        img_float = img.astype('float32') / 255.0
        return img_float

def show_image_version(img_path):
    img = plt.imread(img_path)
    print("Wczytano:", img_path, "-> typ:", img.dtype, "| kształt:", img.shape)

    if img.ndim == 3 and img.shape[-1] == 4:
        img = img[:, :, :3]  # usunięcie kanału alfa

    R = img[:, :, 0]
    G = img[:, :, 1]
    B = img[:, :, 2]

    Y1 = 0.299 * R + 0.587 * G + 0.114 * B
    Y2 = 0.2126 * R + 0.7152 * G + 0.0722 * B

    red_only = img.copy(); red_only[:, :, 1:] = 0
    green_only = img.copy(); green_only[:, :, [0, 2]] = 0
    blue_only = img.copy(); blue_only[:, :, :2] = 0

    plt.figure(figsize=(10, 8))
    plt.suptitle("Zadanie 2 - przetwarzanie obrazu")

    plt.subplot(3, 3, 1); plt.imshow(img); plt.title("Oryginał (RGB)"); plt.axis('off')
    plt.subplot(3, 3, 2); plt.imshow(Y1, cmap='gray'); plt.title("Skala szarości Y1"); plt.axis('off')
    plt.subplot(3, 3, 3); plt.imshow(Y2, cmap='gray'); plt.title("Skala szarości Y2"); plt.axis('off')

    plt.subplot(3, 3, 4); plt.imshow(R, cmap='gray'); plt.title("Warstwa R"); plt.axis('off')
    plt.subplot(3, 3, 5); plt.imshow(G, cmap='gray'); plt.title("Warstwa G"); plt.axis('off')
    plt.subplot(3, 3, 6); plt.imshow(B, cmap='gray'); plt.title("Warstwa B"); plt.axis('off')

    plt.subplot(3, 3, 7); plt.imshow(red_only); plt.title("Kanał czerwony"); plt.axis('off')
    plt.subplot(3, 3, 8); plt.imshow(green_only); plt.title("Kanał zielony"); plt.axis('off')
    plt.subplot(3, 3, 9); plt.imshow(blue_only); plt.title("Kanał niebieski"); plt.axis('off')

    plt.tight_layout()
    out_path = os.path.join(OUT_DIR, f"image_version_{os.path.basename(img_path)}.png")
    plt.savefig(out_path, dpi=150, bbox_inches='tight')
    print("Zapisano wykres do:", out_path)
    plt.close()


def generate_report():
    document = Document()
    document.add_heading('Raport z analizy obrazów', 0)

    # ---------------- Zadanie 1 ----------------
    document.add_heading("Zadanie 1 - Typy danych i konwersje", level=1)
    document.add_paragraph("Pytania i odpowiedzi:")

    qa = [
        ("1. Co oznacza typ zmiennych?",
         "Typ danych określa sposób przechowywania informacji o pikselu."),
        ("2. Jaki jest zakres wartości?",
         "Dla uint8: 0-255, dla float: 0.0-1.0."),
        ("3. Co oznacza ostatni parametr .shape?",
         "Oznacza liczbę kanałów obrazu (np. RGB lub RGBA)."),
        ("4. Czy dla JPG i PNG otrzymujemy te same wyniki?",
         "Nie, PNG często jest w float32, a JPG w uint8."),
        ("5. Różnica między plt.imread() a cv2.imread()?",
         "Matplotlib wczytuje w RGB, a OpenCV w BGR - wymaga zamiany kanałów.")
    ]
    for q, a in qa:
        document.add_paragraph(q, style="List Bullet")
        document.add_paragraph(a)

    document.add_paragraph("\nDodatkowo - analiza obrazu otwartego za pomocą OpenCV:")
    document.add_paragraph(
        "Po otwarciu obrazu przy użyciu funkcji cv2.imread() typ danych to uint8, "
        "a wartości pikseli mieszczą się w zakresie 0-255. "
        "Ostatni parametr .shape nadal oznacza liczbę kanałów, jednak OpenCV wczytuje obraz w kolejności BGR, "
        "czyli kanały kolorów są odwrócone względem RGB."
    )

    document.add_paragraph("\nAnaliza plików A1-A4:")
    table = document.add_table(rows=1, cols=5)
    hdr = table.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text, hdr[3].text, hdr[4].text = "Plik", "Typ danych", "Rozmiar", "Min", "Max"

    for name in ["A1.png", "A2.jpg", "A3.png", "A4.jpg"]:
        path = os.path.join(IMG_DIR, name)
        row = table.add_row().cells
        if os.path.exists(path):
            img = plt.imread(path)
            row[0].text, row[1].text = name, str(img.dtype)
            row[2].text, row[3].text, row[4].text = str(img.shape), str(np.min(img)), str(np.max(img))
        else:
            row[0].text, row[1].text = name, "Brak pliku"

    # ---------------- Zadanie 2 ----------------
    document.add_heading("Zadanie 2 - Wyświetlanie i przetwarzanie obrazu", level=1)
    document.add_paragraph(
        "Obrazy pokazują sposób działania funkcji konwersji do skali szarości oraz separację kanałów RGB."
    )
    document.add_paragraph(
        "Po wczytaniu obrazu przez OpenCV (cv2.imread) kolory są przekłamane, "
        "ponieważ biblioteka używa kolejności kanałów BGR zamiast RGB. "
        "Dlatego obraz wygląda niepoprawnie, dopóki nie zamienimy kanałów funkcją cv2.cvtColor(img, cv2.COLOR_BGR2RGB)."
    )
    document.add_paragraph("Co nam się wyświetliło? - obraz o zmienionych kolorach (nieprawidłowych).")
    document.add_paragraph("Co jest nie tak? - OpenCV zamienia kolejność kanałów, dlatego kolory są odwrócone.")

    for img_name in ["B01.png", "B02.jpg"]:
        img_path = os.path.join(IMG_DIR, img_name)
        if os.path.exists(img_path):
            show_image_version(img_path)
            out_path = os.path.join(OUT_DIR, f"image_version_{os.path.basename(img_path)}.png")
            if os.path.exists(out_path):
                document.add_picture(out_path, width=Inches(6))
                document.add_paragraph(f"Wykres przetwarzania dla: {img_name}")
        else:
            document.add_paragraph(f"Brak pliku: {img_name}")

    # ---------------- Zadanie 3 ----------------
    document.add_heading("Zadanie 3 - Analiza fragmentów obrazów", level=1)
    document.add_paragraph(
        "Wycięte fragmenty przedstawiają różne części tęczy."
    )

    df = pd.DataFrame(data={
        'Filename': ['B01.png', 'B02.jpg'],
        'Grayscale': [False, False],
        'Fragments': [
            [[100, 100, 300, 300], [400, 200, 600, 400]],
            [[150, 150, 350, 350]]
        ]
    })

    for _, row in df.iterrows():
        file_path = os.path.join(IMG_DIR, row['Filename'])
        if not os.path.exists(file_path):
            continue
        document.add_heading(f"Plik: {row['Filename']}", level=2)
        img = plt.imread(file_path)
        for f in row['Fragments']:
            fragment = img[f[0]:f[2], f[1]:f[3]].copy()
            plt.figure(figsize=(6, 4))
            plt.imshow(fragment)
            plt.title(f"Fragment: {f}")
            plt.tight_layout()
            memfile = BytesIO()
            plt.savefig(memfile, format="png")
            memfile.seek(0)
            document.add_picture(memfile, width=Inches(4))
            memfile.close()
            plt.close()


    # Zapis raportu DOCX
    full_path = os.path.join(OUT_DIR, "report.docx")
    document.save(full_path)
    print(f"\nRaport zapisany jako: {full_path}")

    try:
        import subprocess
        subprocess.run(["libreoffice", "--headless", "--convert-to", "pdf", full_path])
        print("Zapisano również wersję PDF.")
    except Exception:
        print("LibreOffice nie znaleziono - pomijam PDF.")


def main():
    print("\n--- Zadanie obraz 1 ---")
    img_test = np.random.rand(100, 100, 3).astype('float32')
    img_uint = imgToUInt8(img_test)
    img_float = imgToFloat(img_uint)

    print("\n--- Zadanie obraz 2 ---")
    sample_img = os.path.join(IMG_DIR, "B01.png")
    if os.path.exists(sample_img):
        show_image_version(sample_img)
    else:
        print("Brak przykładowego obrazu B01.png")

    print("\n--- Zadanie obraz 3 ---")
    generate_report()


if __name__ == "__main__":
    main()