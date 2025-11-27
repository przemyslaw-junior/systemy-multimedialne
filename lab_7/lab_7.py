from __future__ import annotations

import csv
from dataclasses import dataclass, field
from pathlib import Path
from typing import Dict, Iterable, List, Tuple

import cv2
import matplotlib.pyplot as plt
import numpy as np
import scipy.fftpack
from matplotlib import rcParams

try:
    from docx import Document
    from docx.shared import Inches
except Exception:
    Document = None
    Inches = None

BASE_DIR = Path(__file__).parent
RESULTS_DIR = BASE_DIR / "results"
PLOTS_DIR = RESULTS_DIR / "plots"
COMPRESSED_DIR = RESULTS_DIR / "compressed"
TABLE_DIR = RESULTS_DIR / "tables"

HEAD_COLOR = "#1f497d"
BODY_FONT = "DejaVu Sans"
HEAD_SIZE = 18
SUBHEAD_SIZE = 12
BODY_SIZE = 10
FIGSIZE_A4 = (8.27, 11.69)

QY50 = np.array(
    [
        [16, 11, 10, 16, 24, 40, 51, 61],
        [12, 12, 14, 19, 26, 58, 60, 55],
        [14, 13, 16, 24, 40, 57, 69, 56],
        [14, 17, 22, 29, 51, 87, 80, 62],
        [18, 22, 37, 56, 68, 109, 103, 77],
        [24, 36, 55, 64, 81, 104, 113, 92],
        [49, 64, 78, 87, 103, 121, 120, 101],
        [72, 92, 95, 98, 112, 100, 103, 99],
    ]
)
QC50 = np.array(
    [
        [17, 18, 24, 47, 99, 99, 99, 99],
        [18, 21, 26, 66, 99, 99, 99, 99],
        [24, 26, 56, 99, 99, 99, 99, 99],
        [47, 66, 99, 99, 99, 99, 99, 99],
        [99, 99, 99, 99, 99, 99, 99, 99],
        [99, 99, 99, 99, 99, 99, 99, 99],
        [99, 99, 99, 99, 99, 99, 99, 99],
        [99, 99, 99, 99, 99, 99, 99, 99],
    ]
)
Q_NEUTRAL = np.ones((8, 8), dtype=int)

# Container ver2
@dataclass
class JPEGContainer:
    Y: np.ndarray
    Cb: np.ndarray
    Cr: np.ndarray
    shape: Tuple[int, int, int]
    padded_shape_Y: Tuple[int, int]
    padded_shape_C: Tuple[int, int]
    ChromaRatio: str = "4:4:4"
    QY: np.ndarray = field(default_factory=lambda: QY50.copy())
    QC: np.ndarray = field(default_factory=lambda: QC50.copy())
    minY: float = 0.0
    maxY: float = 255.0
    minCb: float = 0.0
    maxCb: float = 255.0
    minCr: float = 0.0
    maxCr: float = 255.0


def ensure_dirs() -> None:
    for d in (RESULTS_DIR, PLOTS_DIR, COMPRESSED_DIR, TABLE_DIR):
        d.mkdir(parents=True, exist_ok=True)


def pad_to_block(arr: np.ndarray, block: int = 8) -> Tuple[np.ndarray, Tuple[int, int]]:
    """Wypełnienie tablicy 2D wielokrotnościami bloku, używając wartości krawędzi."""
    H, W = arr.shape
    new_h = ((H + block - 1) // block) * block
    new_w = ((W + block - 1) // block) * block
    pad_h = new_h - H
    pad_w = new_w - W
    if pad_h == 0 and pad_w == 0:
        return arr, (H, W)
    padded = np.pad(arr, ((0, pad_h), (0, pad_w)), mode="edge")
    return padded, (new_h, new_w)


def restore_range(arr: np.ndarray, vmin: float, vmax: float) -> np.ndarray:
    if vmax == vmin:
        return np.full_like(arr, vmin)
    scaled = (arr - arr.min()) / (arr.max() - arr.min() + 1e-12)
    return scaled * (vmax - vmin) + vmin


def dct2(a: np.ndarray) -> np.ndarray:
    return scipy.fftpack.dct(scipy.fftpack.dct(a.astype(float), axis=0, norm="ortho"), axis=1, norm="ortho")


def idct2(a: np.ndarray) -> np.ndarray:
    return scipy.fftpack.idct(scipy.fftpack.idct(a.astype(float), axis=0, norm="ortho"), axis=1, norm="ortho")


_ZZ_TEMPLATE = np.array(
    [
        [0, 1, 5, 6, 14, 15, 27, 28],
        [2, 4, 7, 13, 16, 26, 29, 42],
        [3, 8, 12, 17, 25, 30, 41, 43],
        [9, 11, 18, 24, 31, 40, 44, 53],
        [10, 19, 23, 32, 39, 45, 52, 54],
        [20, 22, 33, 38, 46, 51, 55, 60],
        [21, 34, 37, 47, 50, 56, 59, 61],
        [35, 36, 48, 49, 57, 58, 62, 63],
    ]
)


def zigzag(block: np.ndarray) -> np.ndarray:
    out = np.empty((64,), dtype=block.dtype)
    for r in range(8):
        for c in range(8):
            out[_ZZ_TEMPLATE[r, c]] = block[r, c]
    return out


def inverse_zigzag(vec: np.ndarray) -> np.ndarray:
    block = np.empty((8, 8), dtype=vec.dtype)
    for r in range(8):
        for c in range(8):
            block[r, c] = vec[_ZZ_TEMPLATE[r, c]]
    return block


def encode_rle(arr: np.ndarray) -> np.ndarray:
    if arr.size == 0:
        return np.array([], dtype=int)
    values = []
    counts = []
    prev = arr[0]
    cnt = 1
    for v in arr[1:]:
        if v == prev:
            cnt += 1
        else:
            values.append(prev)
            counts.append(cnt)
            prev = v
            cnt = 1
    values.append(prev)
    counts.append(cnt)
    return np.vstack((values, counts)).T.flatten().astype(int)


def decode_rle(encoded: np.ndarray) -> np.ndarray:
    if encoded.size == 0:
        return np.array([], dtype=int)
    assert encoded.size % 2 == 0
    pairs = encoded.reshape(-1, 2)
    out = []
    for val, cnt in pairs:
        out.extend([val] * int(cnt))
    return np.array(out, dtype=int)


# Block-level 
def CompressBlock(block: np.ndarray, Q: np.ndarray) -> np.ndarray:
    centered = block.astype(float) - 128.0
    coeffs = dct2(centered)
    quantized = np.round(coeffs / Q).astype(int)
    return zigzag(quantized)


def DecompressBlock(vec: np.ndarray, Q: np.ndarray) -> np.ndarray:
    coeffs = inverse_zigzag(vec).astype(float) * Q
    rec = idct2(coeffs) + 128.0
    return rec


def CompressLayer(L: np.ndarray, Q: np.ndarray) -> np.ndarray:
    S: List[np.ndarray] = []
    for w in range(0, L.shape[0], 8):
        for k in range(0, L.shape[1], 8):
            block = L[w : w + 8, k : k + 8]
            S.append(CompressBlock(block, Q))
    return np.concatenate(S, axis=0)


def DecompressLayer(S: np.ndarray, Q: np.ndarray, shape: Tuple[int, int]) -> np.ndarray:
    H, W = shape
    L = np.zeros((H, W), dtype=float)
    idx = 0
    for w in range(0, H, 8):
        for k in range(0, W, 8):
            vec = S[idx : idx + 64]
            L[w : w + 8, k : k + 8] = DecompressBlock(vec, Q)
            idx += 64
    return L


def chroma_subsample(layer: np.ndarray, ratio: str) -> np.ndarray:
    if ratio == "4:2:2":
        return layer[:, ::2]
    return layer


def chroma_resample(layer: np.ndarray, ratio: str) -> np.ndarray:
    if ratio == "4:2:2":
        return np.repeat(layer, 2, axis=1)
    return layer


# Main API
def CompressJPEG(RGB: np.ndarray, Ratio: str = "4:4:4", QY: np.ndarray = QY50, QC: np.ndarray = QC50) -> JPEGContainer:
    YCrCb = cv2.cvtColor(RGB, cv2.COLOR_RGB2YCrCb).astype(int)
    Y = YCrCb[:, :, 0]
    Cr = YCrCb[:, :, 1]
    Cb = YCrCb[:, :, 2]

    # minima i maksima do przywrócenia dynamiki
    minY, maxY = float(Y.min()), float(Y.max())
    minCb, maxCb = float(Cb.min()), float(Cb.max())
    minCr, maxCr = float(Cr.min()), float(Cr.max())

    Cr_ss = chroma_subsample(Cr, Ratio)
    Cb_ss = chroma_subsample(Cb, Ratio)

    # padding do bloków 8x8
    Y_pad, Y_shape_pad = pad_to_block(Y, 8)
    Cr_pad, C_shape_pad = pad_to_block(Cr_ss, 8)
    Cb_pad, _ = pad_to_block(Cb_ss, 8)

    Y_stream = CompressLayer(Y_pad, QY)
    Cr_stream = CompressLayer(Cr_pad, QC)
    Cb_stream = CompressLayer(Cb_pad, QC)

    Y_rle = encode_rle(Y_stream)
    Cr_rle = encode_rle(Cr_stream)
    Cb_rle = encode_rle(Cb_stream)

    return JPEGContainer(
        Y=Y_rle,
        Cb=Cb_rle,
        Cr=Cr_rle,
        shape=RGB.shape,
        padded_shape_Y=Y_shape_pad,
        padded_shape_C=C_shape_pad,
        ChromaRatio=Ratio,
        QY=QY,
        QC=QC,
        minY=minY,
        maxY=maxY,
        minCb=minCb,
        maxCb=maxCb,
        minCr=minCr,
        maxCr=maxCr,
    )


def DecompressJPEG(JPEG: JPEGContainer) -> np.ndarray:
    H, W, _ = JPEG.shape

    Y_stream = decode_rle(JPEG.Y)
    Cr_stream = decode_rle(JPEG.Cr)
    Cb_stream = decode_rle(JPEG.Cb)

    Y_rec = DecompressLayer(Y_stream, JPEG.QY, JPEG.padded_shape_Y)
    Cr_ss = DecompressLayer(Cr_stream, JPEG.QC, JPEG.padded_shape_C)
    Cb_ss = DecompressLayer(Cb_stream, JPEG.QC, JPEG.padded_shape_C)

    Cr = chroma_resample(Cr_ss, JPEG.ChromaRatio)
    Cb = chroma_resample(Cb_ss, JPEG.ChromaRatio)

    # przywrócenie zakresów i przycięcie do oryginalnego kształtu
    Y_rec = restore_range(Y_rec, JPEG.minY, JPEG.maxY)[: H, : W]
    Cr = restore_range(Cr, JPEG.minCr, JPEG.maxCr)[: H, : W]
    Cb = restore_range(Cb, JPEG.minCb, JPEG.maxCb)[: H, : W]

    YCrCb_rec = np.stack([Y_rec, Cr, Cb], axis=2)
    RGB_rec = cv2.cvtColor(np.clip(YCrCb_rec, 0, 255).astype(np.uint8), cv2.COLOR_YCrCb2RGB)
    return RGB_rec


# Experiment automation
def mse(a: np.ndarray, b: np.ndarray) -> float:
    a = a.astype(float)
    b = b.astype(float)
    return float(np.mean((a - b) ** 2))


def snr_db(a: np.ndarray, b: np.ndarray, eps: float = 1e-12) -> float:
    a = a.astype(float)
    b = b.astype(float)
    num = np.sum(a ** 2)
    den = np.sum((a - b) ** 2) + eps
    return 10 * np.log10(num / den)


def load_images(image_dir: Path) -> List[np.ndarray]:
    imgs = []
    for p in sorted(image_dir.glob("*.png")) + sorted(image_dir.glob("*.jpg")) + sorted(image_dir.glob("*.jpeg")):
        img = cv2.cvtColor(cv2.imread(str(p)), cv2.COLOR_BGR2RGB)
        imgs.append(img)
        if len(imgs) >= 4:
            break
    return imgs


def extract_fragments(img: np.ndarray, size: int = 128, n_frag: int = 3) -> List[np.ndarray]:
    H, W, _ = img.shape
    frags = []
    xs = np.linspace(0, max(0, W - size), n_frag, dtype=int)
    ys = np.linspace(0, max(0, H - size), n_frag, dtype=int)
    for x, y in zip(xs, ys):
        frags.append(img[y : y + size, x : x + size])
    return frags[:n_frag]


def save_comparison_plot(orig: np.ndarray, recon: np.ndarray, title: str, path: Path) -> None:
    fig, axs = plt.subplots(1, 2, figsize=(8, 4))
    axs[0].imshow(orig)
    axs[0].set_title("Original")
    axs[1].imshow(recon)
    axs[1].set_title("Decoded")
    for ax in axs:
        ax.axis("off")
    fig.suptitle(title)
    plt.tight_layout()
    plt.savefig(path, dpi=150)
    plt.close(fig)


def run_experiments() -> List[Dict[str, object]]:
    ensure_dirs()
    image_dir = BASE_DIR / "images"
    images = load_images(image_dir)
    if len(images) == 0:
        print("Nie znaleziono obrazów w katalogu „images”. Dodaj 4 obrazy w wysokiej rozdzielczości.")
        return []
    variants = [
        ("4:4:4", QY50, QC50, "Q50"),
        ("4:4:4", Q_NEUTRAL, Q_NEUTRAL, "Q1"),
        ("4:2:2", QY50, QC50, "Q50"),
        ("4:2:2", Q_NEUTRAL, Q_NEUTRAL, "Q1"),
    ]
    results: List[Dict[str, object]] = []
    for img_idx, img in enumerate(images[:4]):
        frags = extract_fragments(img, size=128, n_frag=3)
        for frag_idx, frag in enumerate(frags):
            for ratio, qy, qc, qname in variants:
                safe_ratio = ratio.replace(":", "-")
                jpeg = CompressJPEG(frag, Ratio=ratio, QY=qy, QC=qc)
                recon = DecompressJPEG(jpeg)
                np.savez_compressed(
                    COMPRESSED_DIR / f"img{img_idx}_frag{frag_idx}_{safe_ratio}_{qname}.npz",
                    Y=jpeg.Y,
                    Cb=jpeg.Cb,
                    Cr=jpeg.Cr,
                    shape=jpeg.shape,
                    ratio=jpeg.ChromaRatio,
                    QY=jpeg.QY,
                    QC=jpeg.QC,
                )
                m = mse(frag, recon)
                s = snr_db(frag, recon)
                results.append(
                    {
                        "image": img_idx,
                        "fragment": frag_idx,
                        "variant": f"{ratio}_{qname}",
                        "ratio": ratio,
                        "Q": qname,
                        "mse": m,
                        "snr": s,
                        "len_Y": jpeg.Y.size,
                        "len_Cb": jpeg.Cb.size,
                        "len_Cr": jpeg.Cr.size,
                        "len_Y_bytes": jpeg.Y.nbytes,
                        "len_Cb_bytes": jpeg.Cb.nbytes,
                        "len_Cr_bytes": jpeg.Cr.nbytes,
                    }
                )
                plot_path = PLOTS_DIR / f"cmp_img{img_idx}_frag{frag_idx}_{safe_ratio}_{qname}.png"
                save_comparison_plot(frag, recon, f"img{img_idx} frag{frag_idx} {ratio} {qname}", plot_path)
    return results


# CSV and report
def save_table(results: List[Dict[str, object]]) -> Path:
    TABLE_DIR.mkdir(parents=True, exist_ok=True)
    path = TABLE_DIR / "table.csv"
    if not results:
        return path
    fieldnames = [
        "image",
        "fragment",
        "variant",
        "ratio",
        "Q",
        "mse",
        "snr",
        "len_Y",
        "len_Cb",
        "len_Cr",
        "len_Y_bytes",
        "len_Cb_bytes",
        "len_Cr_bytes",
    ]
    with path.open("w", newline="", encoding="utf-8") as f:
        writer = csv.DictWriter(f, fieldnames=fieldnames)
        writer.writeheader()
        writer.writerows(results)
    return path


def build_report_docx(results: List[Dict[str, object]]) -> None:
    if Document is None:
        print("[INFO] python-docx niedostępne - pomijam raport DOCX.")
        return
    doc = Document()
    doc.add_heading("Laboratorium 7 — Uproszczony JPEG", level=0)
    doc.add_paragraph(
        "Celem jest implementacja uproszczonego JPEG: RGB->YCrCb, redukcja chrominancji 4:2:2, bloki 8x8, "
        "DCT/IDCT, kwantyzacja (QY/QC lub neutralna), zygzakowanie i RLE. "
        "Porównano 4 warianty dla fragmentów obrazów 128x128."
    )

    doc.add_heading("Przetwarzanie", level=1)
    doc.add_paragraph(
        "Pipeline: konwersja RGB->YCrCb, subsampling 4:4:4 lub 4:2:2, DCT 8x8, kwantyzacja Q50 lub Q=1, "
        "zygzakowanie, RLE per warstwa, dekwantyzacja i IDCT."
    )

    doc.add_heading("Porównania wizualne", level=1)
    for plot_path in sorted(PLOTS_DIR.glob("cmp_*.png")):
        if plot_path.exists() and Inches:
            doc.add_picture(str(plot_path), width=Inches(6))

    doc.add_heading("Wyniki — metryki", level=1)
    if results:
        t = doc.add_table(rows=1, cols=9)
        h = t.rows[0].cells
        h[0].text = "Obraz"
        h[1].text = "Frag"
        h[2].text = "Wariant"
        h[3].text = "SNR [dB]"
        h[4].text = "MSE"
        h[5].text = "len(Y)"
        h[6].text = "len(Cb)"
        h[7].text = "len(Cr)"
        h[8].text = "Bajty Y/Cb/Cr"
        for r in results:
            row = t.add_row().cells
            row[0].text = str(r["image"])
            row[1].text = str(r["fragment"])
            row[2].text = str(r["variant"])
            row[3].text = f"{r['snr']:.2f}"
            row[4].text = f"{r['mse']:.2e}"
            row[5].text = str(r["len_Y"])
            row[6].text = str(r["len_Cb"])
            row[7].text = str(r["len_Cr"])
            row[8].text = f"{r['len_Y_bytes']}/{r['len_Cb_bytes']}/{r['len_Cr_bytes']}"

    doc.add_heading("Wnioski", level=1)
    doc.add_paragraph(
        "Subsampling 4:2:2 i tablice Q50 zwiększają kompresję kosztem artefaktów na chrominancji. "
        "Warianty Q=1 minimalizują zniekształcenia, ale kompresja jest mniejsza. "
        "Fragmenty o gładkich obszarach kompresują się lepiej niż te z drobną teksturą."
    )

    path = RESULTS_DIR / "report.docx"
    doc.save(str(path))
    print(f"[OK] Zapisano raport DOCX: {path}")


def main() -> None:
    ensure_dirs()
    results = run_experiments()
    save_table(results)
    build_report_docx(results)


if __name__ == "__main__":
    main()
    