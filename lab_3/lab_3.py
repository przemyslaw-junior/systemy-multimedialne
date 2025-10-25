import os
import numpy as np
import matplotlib.pyplot as plt
import cv2
import subprocess
import shutil
from docx import Document
from docx.shared import Inches

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
IMG_DIR = os.path.join(BASE_DIR, "images")
OUT_DIR = os.path.join(BASE_DIR, "output")
os.makedirs(IMG_DIR, exist_ok=True)
os.makedirs(OUT_DIR, exist_ok=True)

def to_float01(img):
    if img.dtype.kind in ("u", "i"):
        return img.astype("float32") / 255.0
    if img.dtype.kind == "f":
        if np.nanmax(img) > 1.0 or np.nanmin(img) < 0.0:
            return np.clip(img, 0, 255).astype("float32") / 255.0
        return img.astype("float32")
    return to_float01(img.astype("float32"))

def load_gray(path):
    g = cv2.imread(path, cv2.IMREAD_GRAYSCALE)
    if g is None:
        raise FileNotFoundError(path)
    g = to_float01(g)
    return g[..., None]

def load_rgb(path):
    bgr = cv2.imread(path, cv2.IMREAD_COLOR)
    if bgr is None:
        raise FileNotFoundError(path)
    rgb = cv2.cvtColor(bgr, cv2.COLOR_BGR2RGB)
    return to_float01(rgb)

def save_image(fname, arr):
    path = os.path.join(OUT_DIR, fname)
    arr = np.clip(arr, 0.0, 1.0)
    if arr.ndim == 3 and arr.shape[2] == 1:
        plt.imsave(path, arr[:, :, 0], cmap="gray", vmin=0, vmax=1)
    else:
        plt.imsave(path, arr)
    print("Zapisano:", path)
    return path

def mosaic2x3(title, images, labels, out_name):
    plt.figure(figsize=(12, 8))
    plt.suptitle(title)
    for i, (img, lab) in enumerate(zip(images, labels), 1):
        plt.subplot(2, 3, i)
        if img.ndim == 3 and img.shape[2] == 1:
            plt.imshow(img[:, :, 0], cmap="gray", vmin=0, vmax=1)
        else:
            plt.imshow(img)
        plt.title(lab); plt.axis("off")
    plt.tight_layout(rect=[0, 0, 1, 0.96])
    path = os.path.join(OUT_DIR, out_name)
    plt.savefig(path, dpi=150, bbox_inches="tight")
    plt.close()
    print("Zapisano mozaikę:", path)
    return path

def gray_palette(levels):
    return np.linspace(0.0, 1.0, levels, dtype="float32").reshape(-1, 1)

pallet8 = np.array([
    [0.0, 0.0, 0.0],
    [0.0, 0.0, 1.0],
    [0.0, 1.0, 0.0],
    [0.0, 1.0, 1.0],
    [1.0, 0.0, 0.0],
    [1.0, 0.0, 1.0],
    [1.0, 1.0, 0.0],
    [1.0, 1.0, 1.0],
], dtype="float32")

pallet16 = np.array([
    [0.0, 0.0, 0.0],
    [0.0, 1.0, 1.0],
    [0.0, 0.0, 1.0],
    [1.0, 0.0, 1.0],
    [0.0, 0.5, 0.0],
    [0.5, 0.5, 0.5],
    [0.0, 1.0, 0.0],
    [0.5, 0.0, 0.0],
    [0.0, 0.0, 0.5],
    [0.5, 0.5, 0.0],
    [0.5, 0.0, 0.5],
    [1.0, 0.0, 0.0],
    [0.75, 0.75, 0.75],
    [0.0, 0.5, 0.5],
    [1.0, 1.0, 1.0],
    [1.0, 1.0, 0.0],
], dtype="float32")

def colorFit(pixel, Pallet):
    px = np.array(pixel, dtype="float32").reshape(1, -1)
    diffs = Pallet - px  # N x C
    dists = np.linalg.norm(diffs, axis=1)  # N
    idx = int(np.argmin(dists))
    return Pallet[idx]

def kwant_colorFit(img, Pallet):
    H, W, C = img.shape
    out = np.empty_like(img, dtype="float32")
    for w in range(H):
        for k in range(W):
            out[w, k] = colorFit(img[w, k], Pallet)
    return out

def dither_random(img_gray_01, Pallet):
    assert img_gray_01.shape[2] == 1, "dither_random: wymagany obraz 1-warstwowy"
    assert Pallet.shape[0] == 2 and Pallet.shape[1] == 1, "paleta musi mieć 2 poziomy (binarny)"
    H, W, _ = img_gray_01.shape
    r = np.random.rand(H, W, 1).astype("float32")
    mask = (img_gray_01 >= r).astype("float32")
    lo, hi = Pallet[0, 0], Pallet[1, 0]
    return lo * (1.0 - mask) + hi * mask

def bayer_matrix_4x4():
    return np.array([
        [0,  8,  2, 10],
        [12, 4, 14,  6],
        [3, 11,  1,  9],
        [15, 7, 13,  5],
    ], dtype="float32")

def dither_ordered(img, Pallet, r=1.0):
    H, W, C = img.shape
    M = bayer_matrix_4x4()
    n = 4.0
    Mpre = (M + 0.5) / (n * n) - 0.5
    tile = np.tile(Mpre, (int(np.ceil(H / n)), int(np.ceil(W / n))))
    tile = tile[:H, :W].astype("float32")
    tileC = np.repeat(tile[:, :, None], C, axis=2)
    tmp = np.clip(img + r * tileC, 0.0, 1.0)
    out = np.empty_like(tmp)
    for w in range(H):
        for k in range(W):
            out[w, k] = colorFit(tmp[w, k], Pallet)
    return out

def dither_floyd_steinberg(img, Pallet):
    H, W, C = img.shape
    out = img.copy().astype("float32")
    for y in range(H):
        for x in range(W):
            old = out[y, x].copy()
            new = colorFit(old, Pallet)
            out[y, x] = new
            err = old - new
            if x + 1 < W:
                out[y, x + 1] += err * (7.0 / 16.0)
            if y + 1 < H and x - 1 >= 0:
                out[y + 1, x - 1] += err * (3.0 / 16.0)
            if y + 1 < H:
                out[y + 1, x] += err * (5.0 / 16.0)
            if y + 1 < H and x + 1 < W:
                out[y + 1, x + 1] += err * (1.0 / 16.0)
    return np.clip(out, 0.0, 1.0)

def add_picture_if_exists(doc, path, width_in=6.0, caption=None):
    if os.path.exists(path):
        doc.add_picture(path, width=Inches(width_in))
        if caption:
            doc.add_paragraph(caption)


def build_report(records, conclusions_short: str = ""):
    document = Document()
    document.add_heading("Kwantyzacja obrazu i dithering", 0)

    # Zadanie 1 – colorFit i palety
    document.add_heading("Zadanie 1 - colorFit i palety", level=1)
    document.add_paragraph("Funkcja colorFit dopasowuje kolor piksela do najbliższego koloru z wybranej palety.")
    document.add_paragraph("Działa zarówno dla obrazów w skali szarości (Nx1), jak i kolorowych (Nx3), wybierając kolor o najmniejszej różnicy (odległości) w przestrzeni barw.")

    # Zadanie 2 – Kwantyzacja
    document.add_heading("Zadanie 2 - Kwantyzacja", level=1)
    document.add_paragraph("Kwantyzacja polega na zmniejszeniu liczby poziomów jasności lub kolorów w obrazie.")
    document.add_paragraph("Dla obrazów szarych wykonano redukcję do 1, 2 i 4 bitów (2, 4 i 16 odcieni).")
    document.add_paragraph("Dla obrazów kolorowych użyto gotowych palet z 8 i 16 kolorami.")
    document.add_paragraph("Po redukcji widać wyraźnie utratę szczegółów i pojawianie się pasm kolorów (tzw. banding).")

    # Zadanie 3 – Dithering
    document.add_heading("Zadanie 3 - Dithering", level=1)
    document.add_paragraph("Zaimplementowano trzy metody ditheringu:")
    document.add_paragraph("Losowy - działa tylko dla obrazów binarnych (czarno-białych), dodaje szum, ale poprawia widoczność półtonów.")
    document.add_paragraph("Zorganizowany (Bayer 4x4) - tworzy równy wzór punktowy, daje lepszy efekt niż metoda losowa.")
    document.add_paragraph("Floyd-Steinberg - najlepsza metoda; rozprasza błędy kwantyzacji na sąsiednie piksele, dzięki czemu obraz wygląda bardziej naturalnie i szczegółowo.")
    document.add_paragraph("Dithering pomaga ukryć utratę jakości po kwantyzacji i daje wrażenie większej liczby kolorów.")
  
    document.add_heading("Wnioski", level=1)
    if conclusions_short.strip() == "":
        conclusions = [
            "Kwantyzacja znacząco obniża jakość obrazu, szczególnie przy małej liczbie bitów.",
            "Zastosowanie ditheringu poprawia wygląd obrazu i pozwala uzyskać lepsze przejścia tonalne.",
            "Najlepsze rezultaty daje metoda Floyd-Steinberg, natomiast zorganizowany Bayer jest szybki i przewidywalny, a losowy daje najprostszy, ale najmniej dokładny efekt.",
        ]
        for p in conclusions:
            document.add_paragraph(p)
    else:
        document.add_paragraph(conclusions_short)

  
    document.add_heading("Wyniki i przykłady", level=1)
    for title, p in records:
        add_picture_if_exists(document, p, width_in=6.0, caption=title)

    path = os.path.join(BASE_DIR, "report.docx")
    document.save(path)
    print("Zapisano raport:", path)

  
    try:
        try:
            from docx2pdf import convert
            out_pdf = os.path.splitext(path)[0] + ".pdf"
            convert(path, out_pdf)
            if os.path.exists(out_pdf):
                print("Zapisano również wersję PDF:", out_pdf)
                return
        except Exception:
            pass
        for exe in ("soffice", "libreoffice", "lowriter"):
            if shutil.which(exe):
                try:
                    subprocess.run([exe, "--headless", "--convert-to", "pdf", path, "--outdir", BASE_DIR], check=True)
                    out_pdf = os.path.splitext(path)[0] + ".pdf"
                    if os.path.exists(out_pdf):
                        print("Zapisano również wersję PDF:", out_pdf)
                    else:
                        print("Konwersja do PDF wykonana, ale plik nieodnaleziony.")
                    return
                except Exception:
                    continue
        print("Konwersja do PDF pominięta (brak docx2pdf/LibreOffice).")
    except Exception:
        print("Konwersja do PDF nieudana (wyjątek).")

def main():
    gs_files = ["GS_0001.tif", "GS_0002.png", "GS_0003.png"]
    color_files = ["SMALL_0001.tif", "SMALL_0004.jpg", "SMALL_0005.jpg", "SMALL_0006.jpg"]

    report_items = []

    for name in gs_files:
        path = os.path.join(IMG_DIR, name)
        if not os.path.exists(path):
            print("Brak pliku:", path)
            continue

        g = load_gray(path)

        q1 = kwant_colorFit(g, gray_palette(2))
        q2 = kwant_colorFit(g, gray_palette(4))
        q4 = kwant_colorFit(g, gray_palette(16))

        pal2 = gray_palette(2)
        d_rand = dither_random(g, pal2)
        d_ord = dither_ordered(g, pal2, r=1.0)
        d_fs = dither_floyd_steinberg(g, pal2)

        title = f"GS: {name}"
        mosaic_path = mosaic2x3(
            title,
            images=[g, q1, q2, q4, d_rand, d_fs],
            labels=["Oryginał", "Kwant 1 bit", "Kwant 2 bity", "Kwant 4 bity", "Dither losowy (bin)", "Floyd-Steinberg (bin)"],
            out_name=f"gs_mosaic_{os.path.splitext(name)[0]}.png",
        )
        report_items.append((title, mosaic_path))

        p2 = save_image(f"gs_ordered_{os.path.splitext(name)[0]}.png", d_ord)
        report_items.append((f"{name} - Dither zorganizowany (Bayer 4x4, r=1)", p2))

        # Dodatkowe przypadki: dithering 2 i 4 bitów (4 i 16 poziomów)
        pal4 = gray_palette(4)
        pal16 = gray_palette(16)
        d_ord4 = dither_ordered(g, pal4, r=1.0)
        d_fs4 = dither_floyd_steinberg(g, pal4)
        d_ord16 = dither_ordered(g, pal16, r=1.0)
        d_fs16 = dither_floyd_steinberg(g, pal16)

        p_ord4 = save_image(f"gs_ordered_2bit_{os.path.splitext(name)[0]}.png", d_ord4)
        report_items.append((f"{name} - Dither zorganizowany (4 poziomy)", p_ord4))
        p_fs4 = save_image(f"gs_fs_2bit_{os.path.splitext(name)[0]}.png", d_fs4)
        report_items.append((f"{name} - Floyd-Steinberg (4 poziomy)", p_fs4))

        p_ord16 = save_image(f"gs_ordered_4bit_{os.path.splitext(name)[0]}.png", d_ord16)
        report_items.append((f"{name} - Dither zorganizowany (16 poziomów)", p_ord16))
        p_fs16 = save_image(f"gs_fs_4bit_{os.path.splitext(name)[0]}.png", d_fs16)
        report_items.append((f"{name} - Floyd-Steinberg (16 poziomów)", p_fs16))

    for name in color_files:
        path = os.path.join(IMG_DIR, name)
        if not os.path.exists(path):
            print("Brak pliku:", path)
            continue

        rgb = load_rgb(path)

        q8 = kwant_colorFit(rgb, pallet8)
        q16 = kwant_colorFit(rgb, pallet16)

        d8_ord = dither_ordered(rgb, pallet8, r=1.0)
        d8_fs = dither_floyd_steinberg(rgb, pallet8)
        
        # Dodatkowe przypadki: dithering dla palety 16 kolorów
        d16_ord = dither_ordered(rgb, pallet16, r=1.0)
        d16_fs = dither_floyd_steinberg(rgb, pallet16)

        title = f"RGB: {name}"
        mosaic_path = mosaic2x3(
            title,
            images=[rgb, q8, q16, d8_ord, d8_fs, q8],
            labels=["Oryginał", "Kwant 8 kolorów", "Kwant 16 kolorów", "Ordered (8 kol.)", "Floyd-Steinberg (8 kol.)", "—"],
            out_name=f"rgb_mosaic_{os.path.splitext(name)[0]}.png",
        )
        report_items.append((title, mosaic_path))

        p16o = save_image(f"rgb_ordered_16col_{os.path.splitext(name)[0]}.png", d16_ord)
        report_items.append((f"{name} - Dither zorganizowany (16 kolorów)", p16o))

        p16f = save_image(f"rgb_fs_16col_{os.path.splitext(name)[0]}.png", d16_fs)
        report_items.append((f"{name} - Floyd-Steinberg (16 kolorów)", p16f))

    build_report(report_items)

    print("Gotowe.")

if __name__ == "__main__":
    main()