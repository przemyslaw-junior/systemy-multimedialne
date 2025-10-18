import os
import numpy as np
import matplotlib
#matplotlib.use('Agg')
import matplotlib.pyplot as plt
import soundfile as sf
import sounddevice as sd
import scipy as sp
from scipy.fftpack import fft, fftfreq
from docx import Document
from docx.shared import Inches
from io import BytesIO

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
PLOT_DIR = os.path.join(BASE_DIR, "plot")
AUDIO_DIR = os.path.join(BASE_DIR, "audio")
os.makedirs(AUDIO_DIR, exist_ok=True)
os.makedirs(PLOT_DIR, exist_ok=True)

def resolve_path(path: str) -> str:
    if os.path.isabs(path):
        return path

    candidate = os.path.join(AUDIO_DIR, path)
    return candidate if os.path.exists(candidate) else os.path.join(BASE_DIR, path)

# zadanie dźwięk 1 - wczytanie, zapis i odtwarzanie audio

# wczytanie pliku audio
def load_audio_file(path):
    full_path = resolve_path(path)
    if not os.path.exists(full_path):
        raise FileNotFoundError(f"nie znaleziono pliku: {full_path}")
    data, fs = sf.read(full_path, dtype='float32')
    print("wczytano plik z:", full_path)
    print("typ danych:", data.dtype, "| kształt:", getattr(data, 'shape', 'brak'))
    return data, fs

# zapisanie pliku audio + kanały
def save_audio_file(data, fs, out_dir=None):
    if out_dir is None:
        out_dir = AUDIO_DIR
    os.makedirs(out_dir, exist_ok=True)
    if data.ndim < 2:
        print("zapis mono")
        sound_L = sound_R = data
    else:
        print("zapis stereo")
        sound_L = data[:, 0]
        sound_R = data[:, 1]
    sound_mix = ((sound_L + sound_R) / 2).astype(np.float32)

    sf.write(os.path.join(out_dir, 'sound_L.wav'), sound_L, fs)
    sf.write(os.path.join(out_dir, 'sound_R.wav'), sound_R, fs)
    sf.write(os.path.join(out_dir, 'sound_mix.wav'), sound_mix, fs)

    print("zapisano pliki do:", out_dir)

# odtworzenie pliku audio + czekanie na zakończenie
def play_audio(data, fs, title="odtwarzanie audio"):
    print(title)
    sd.play(data, fs)
    sd.wait()
    print("odtwarzanie zakończone")


# 2. zadanie dźwięk 2 - rysowanie wykresu sygnału audio
def plot_audio(Signal, Fs, TimeMargin=[0, 0.02], save_name=None, out_dir=None):
    plt.figure()

    plt.subplot(2, 1, 1)
    sig_mono = Signal.mean(axis=1) if getattr(Signal, 'ndim', 1) > 1 else Signal
    t = np.arange(0, sig_mono.shape[0]) / Fs
    plt.plot(t, sig_mono)
    plt.xlim(TimeMargin)
    plt.xlabel('Czas [s]')
    plt.ylabel('Amplituda')
    plt.title('Sygnał audio w dziedzinie czasu')

    plt.subplot(2, 1, 2)
    fsize = 2 ** 8
    yf = sp.fftpack.fft(sig_mono, fsize)
    # Wyznaczamy częstotliwości osi X dla widma FFT
    f = np.arange(0, Fs / 2, Fs / fsize)
    plt.plot(f, 20 * np.log10(np.abs(yf[0:fsize // 2]) + 1e-12))
    plt.xlabel('Częstotliwość [Hz]')
    plt.ylabel('Amplituda [dB]')
    plt.title('Widmo sygnału audio')
    plt.tight_layout()
   
    if save_name is not None:
        target_dir = (PLOT_DIR if out_dir is None else out_dir)
        out_path = os.path.join(target_dir, f"{save_name}.png")
        plt.savefig(out_path, dpi=150, bbox_inches='tight')
        print(f"zapisano wykres: {out_path}")
    plt.close()


# 3. zadanie dźwięk 3 - analiza FFT i raport DOCX
def analyze_and_generate_raport(files, fsize_list, out_dir='report.docx'):
    document = Document()
    document.add_heading('Raport z analizy sygnałów audio', 0)
    document.add_heading('Wnioski:', level=1)
    document.add_paragraph(
        "Przeprowadzone zadania potwierdziły poprawne działanie funkcji przetwarzania i analizy sygnałów audio. "
        "Widma uzyskanych przebiegów sinusoidalnych pokazały zgodność częstotliwości z wartościami teoretycznymi, "
        "co potwierdza prawidłowe wykorzystanie transformaty Fouriera. "
    )

    for file in files:
        full_path = resolve_path(file)
        if not os.path.exists(full_path):
            print(f"Plik {full_path} nie istnieje, pomijam")
            continue

        document.add_heading(f'Plik: {full_path}', 1)
        data, fs = sf.read(full_path, dtype='float32')

        for fsize in fsize_list:
            plt.figure()
            plt.subplot(2, 1, 1)
            data_mono = data.mean(axis=1) if getattr(data, 'ndim', 1) > 1 else data
            t = np.arange(0, data_mono.shape[0]) / fs
            plt.plot(t, data_mono)
            plt.xlabel('Czas [s]')
            plt.ylabel('Amplituda')
            plt.title(f'Sygnał audio w dziedzinie czasu (fsize={fsize})')

            plt.subplot(2, 1, 2)
            yf = sp.fftpack.fft(data_mono, fsize)
            f = np.arange(0, fs / 2, fs / fsize)
            magnitude = 20 * np.log10(np.abs(yf[0:fsize // 2]) + 1e-12)
            plt.plot(f, magnitude)
            plt.xlabel('Częstotliwość [Hz]')
            plt.ylabel('Amplituda [dB]')
            plt.title(f'Widmo sygnału (fsize={fsize})')

            max_idx = int(np.argmax(magnitude))
            max_freq = f[max_idx] if len(f) > max_idx else 0.0
            max_amp = magnitude[max_idx] if len(magnitude) > max_idx else 0.0
            plt.title(f'Widmo sygnału (fsize={fsize}) - Max: {max_freq:.2f} Hz, {max_amp:.2f} dB')
            plt.tight_layout()

            stem = os.path.splitext(os.path.basename(full_path))[0]
            fig_name = f"fft_{stem}_fsize{fsize}.png"
            fig_path = os.path.join(PLOT_DIR, fig_name)
            plt.savefig(fig_path, dpi=150, bbox_inches='tight')
            print(f"zapisano wykres: {fig_path}")

            memfile = BytesIO()
            plt.savefig(memfile, format='png')
            memfile.seek(0)
            document.add_heading(f'Analiza FFT (fsize={fsize})', 2)
            document.add_picture(memfile, width=Inches(6))
            document.add_paragraph(f"maksimum widma : {max_freq:.2f} Hz, amplituda: {max_amp:.2f} dB")
            memfile.close()
            plt.close()

    raport_path = resolve_path(out_dir)
    document.save(raport_path)
    print(f'raport zapisany do {raport_path}')


def generate_sine_wav(filename: str, freq: float, duration: float = 2.0, fs: int = 44100, amplitude: float = 0.5):
    t = np.arange(int(duration * fs)) / fs
    y = (amplitude * np.sin(2 * np.pi * freq * t)).astype(np.float32)
    out_path = os.path.join(AUDIO_DIR, filename)
    sf.write(out_path, y, fs)
    print(f"wygenerowano: {out_path} ({freq} Hz)")


def sinus_files():
    targets = [(60, 'sin60Hz.wav'), (440, 'sin440Hz.wav'), (1000, 'sin1000Hz.wav')]
    for freq, name in targets:
        if not os.path.exists(resolve_path(name)):
            generate_sine_wav(name, freq)


def main():
    print("\n--- Zadanie Dźwięk 1 ---")
    try:
        data, fs = load_audio_file('sound1.wav')
        play_audio(data, fs, title="sound1.wav (oryginalny)")
        save_audio_file(data, fs)
    except Exception as e:
        print("błąd w zadaniu 1:", e)
        return

    print("\n--- Zadanie Dźwięk 2 ---")
    sinus_files()
    try:
        sin440, fs_sin = load_audio_file('sin440Hz.wav')
        plot_audio(sin440, fs_sin, save_name='plot_sin440_time_fft')
    except Exception as e:
        print("pomijam zadanie 2:", e)

    print("\n--- Zadanie Dźwięk 3 ---")
    files = ['sin60Hz.wav', 'sin440Hz.wav', 'sin8000Hz.wav']
    file_list = [2 ** 8, 2 ** 12, 2 ** 16]
    analyze_and_generate_raport(files, file_list)


if __name__ == "__main__":
    main()
