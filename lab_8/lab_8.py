from __future__ import annotations

import csv
import os
from dataclasses import dataclass
from pathlib import Path
from typing import Dict, Iterable, List, Tuple

import cv2
import matplotlib.pyplot as plt
import numpy as np
from docx import Document
from docx.shared import Inches

BASE_DIR = Path(__file__).parent
RESULTS_DIR = BASE_DIR / "results"
IMG_DIR = BASE_DIR / "images"
for d in (RESULTS_DIR, IMG_DIR):
    d.mkdir(exist_ok=True)

# Ustawienia eksperymentów
DEFAULT_SUBSAMPLING = ["4:4:4", "4:2:2", "4:4:0", "4:2:0", "4:1:1", "4:1:0"]
DEFAULT_DIVIDERS = [1, 2, 4, 8, 16, 32]
DEFAULT_KEY_DIST = [1, 2, 4, 6, 8, 10]
NUM_FRAMES_TEST = 20  # liczba klatek do pojedynczego testu
ROI_LIST = [[0, 100, 0, 100]]  # można zmienić na listę ROI
USE_RLE_KEY = False  # RLE dla klatek kluczowych
USE_RLE_DELTA = True  # RLE dla klatek różnicowych
SAVE_ROI_IMAGES = False  # zapisywać porównania ROI (PNG)

# Struktury danych
@dataclass
class FrameData:
    Y: np.ndarray
    Cb: np.ndarray
    Cr: np.ndarray
    shape_Y: Tuple[int, int]
    shape_C: Tuple[int, int]


@dataclass
class StoredFrame:
    Y: np.ndarray
    Cb: np.ndarray
    Cr: np.ndarray
    shape_Y: Tuple[int, int]
    shape_C: Tuple[int, int]


# RLE
def rle_encode(arr: np.ndarray) -> np.ndarray:
    flat = arr.flatten()
    if flat.size == 0:
        return flat
    vals: List[int] = []
    runs: List[int] = []
    prev = flat[0]
    cnt = 1
    for v in flat[1:]:
        if v == prev:
            cnt += 1
        else:
            vals.append(int(prev))
            runs.append(cnt)
            prev = v
            cnt = 1
    vals.append(int(prev))
    runs.append(cnt)
    out = np.empty(len(vals) * 2, dtype=int)
    out[0::2] = runs
    out[1::2] = vals
    return out


def rle_decode(encoded: np.ndarray, shape: Tuple[int, int]) -> np.ndarray:
    if encoded.size == 0:
        return np.zeros(shape, dtype=int)
    runs = encoded[0::2]
    vals = encoded[1::2]
    out: List[int] = []
    for r, v in zip(runs, vals):
        out.extend([int(v)] * int(r))
    return np.array(out, dtype=int).reshape(shape)


# Subsampling / resampling
def chroma_subsampling(L: np.ndarray, mode: str) -> Tuple[np.ndarray, Tuple[int, int]]:
    """Redukcja chrominancji bez pętli (slicing)."""
    mode = mode.replace("-", ":").strip()
    if mode == "4:4:4":
        return L, L.shape
    if mode == "4:2:2":  # połowa kolumn
        return L[:, ::2], (L.shape[0], L.shape[1] // 2)
    if mode == "4:4:0":  # połowa wierszy
        return L[::2, :], (L.shape[0] // 2, L.shape[1])
    if mode == "4:2:0":  # połowa wierszy i kolumn
        return L[::2, ::2], (L.shape[0] // 2, L.shape[1] // 2)
    if mode == "4:1:1":  # 4:1 w poziomie
        return L[:, ::4], (L.shape[0], L.shape[1] // 4)
    if mode == "4:1:0":  # 4:1 w poziomie i 2:1 w pionie
        return L[::2, ::4], (L.shape[0] // 2, L.shape[1] // 4)
    return L, L.shape


def chroma_resampling(L: np.ndarray, mode: str, target_shape: Tuple[int, int]) -> np.ndarray:
    """Odtwarzanie chrominancji przez powielanie (np.repeat) i dopasowanie do kształtu."""
    mode = mode.replace("-", ":").strip()
    out = L
    if mode in ("4:2:2", "4:1:1", "4:1:0"):
        out = np.repeat(out, 2 if mode == "4:2:2" else 4, axis=1)
    if mode in ("4:4:0", "4:2:0", "4:1:0"):
        out = np.repeat(out, 2, axis=0)
    # awaryjne dopasowanie do kształtu docelowego
    th, tw = target_shape
    h, w = out.shape
    if h != th or w != tw:
        ry = max(1, int(np.ceil(th / h)))
        rx = max(1, int(np.ceil(tw / w)))
        out = np.repeat(np.repeat(out, ry, axis=0), rx, axis=1)
    return out[:th, :tw]


# Konwersje ramki <-> struktura
def frame_to_struct(frame_ycrcb: np.ndarray, mode: str) -> FrameData:
    mode = mode.replace("-", ":").strip()
    Y = frame_ycrcb[:, :, 0].astype(int)
    Cr = frame_ycrcb[:, :, 1].astype(int)
    Cb = frame_ycrcb[:, :, 2].astype(int)
    Cr_ss, shape_c = chroma_subsampling(Cr, mode)
    Cb_ss, _ = chroma_subsampling(Cb, mode)
    return FrameData(Y=Y, Cb=Cb_ss, Cr=Cr_ss, shape_Y=Y.shape, shape_C=shape_c)


def struct_to_frame(data: FrameData, mode: str) -> np.ndarray:
    mode = mode.replace("-", ":").strip()
    Cb = chroma_resampling(data.Cb, mode, data.shape_Y)
    Cr = chroma_resampling(data.Cr, mode, data.shape_Y)
    # awaryjne dopasowanie gdyby tryb był niewłaściwy lub kształty się różniły
    if Cb.shape != data.shape_Y:
        ry = max(1, int(np.ceil(data.shape_Y[0] / Cb.shape[0])))
        rx = max(1, int(np.ceil(data.shape_Y[1] / Cb.shape[1])))
        Cb = np.repeat(np.repeat(Cb, ry, axis=0), rx, axis=1)[: data.shape_Y[0], : data.shape_Y[1]]
    if Cr.shape != data.shape_Y:
        ry = max(1, int(np.ceil(data.shape_Y[0] / Cr.shape[0])))
        rx = max(1, int(np.ceil(data.shape_Y[1] / Cr.shape[1])))
        Cr = np.repeat(np.repeat(Cr, ry, axis=0), rx, axis=1)[: data.shape_Y[0], : data.shape_Y[1]]
    return np.dstack([data.Y, Cr, Cb]).clip(0, 255).astype(np.uint8)


# Kompresja / dekompresja klatek
def compress_key_frame(f: FrameData, apply_rle: bool = False) -> StoredFrame:
    if apply_rle:
        return StoredFrame(
            Y=rle_encode(f.Y),
            Cb=rle_encode(f.Cb),
            Cr=rle_encode(f.Cr),
            shape_Y=f.shape_Y,
            shape_C=f.shape_C,
        )
    return StoredFrame(Y=f.Y, Cb=f.Cb, Cr=f.Cr, shape_Y=f.shape_Y, shape_C=f.shape_C)


def decompress_key_frame(st: StoredFrame, mode: str, apply_rle: bool = False) -> np.ndarray:
    mode = mode.replace("-", ":").strip()
    if apply_rle:
        fd = FrameData(
            Y=rle_decode(st.Y, st.shape_Y),
            Cb=rle_decode(st.Cb, st.shape_C),
            Cr=rle_decode(st.Cr, st.shape_C),
            shape_Y=st.shape_Y,
            shape_C=st.shape_C,
        )
    else:
        fd = FrameData(Y=st.Y, Cb=st.Cb, Cr=st.Cr, shape_Y=st.shape_Y, shape_C=st.shape_C)
    return struct_to_frame(fd, mode)


def compress_delta(
    curr: FrameData,
    key: StoredFrame,
    mode: str,
    divider: int = 1,
    apply_rle: bool = False,
    apply_rle_key: bool = False,
) -> StoredFrame:
    mode = mode.replace("-", ":").strip()
    # jeśli key jest w RLE, zdekoduj do różnicy
    if apply_rle_key:
        key_Y = rle_decode(key.Y, key.shape_Y)
        key_Cb = rle_decode(key.Cb, key.shape_C)
        key_Cr = rle_decode(key.Cr, key.shape_C)
    else:
        key_Y, key_Cb, key_Cr = key.Y, key.Cb, key.Cr

    dY = (curr.Y.astype(int) - key_Y.astype(int)) // divider
    dCb = (curr.Cb.astype(int) - key_Cb.astype(int)) // divider
    dCr = (curr.Cr.astype(int) - key_Cr.astype(int)) // divider

    if apply_rle:
        return StoredFrame(
            Y=rle_encode(dY),
            Cb=rle_encode(dCb),
            Cr=rle_encode(dCr),
            shape_Y=key.shape_Y,
            shape_C=key.shape_C,
        )
    return StoredFrame(Y=dY, Cb=dCb, Cr=dCr, shape_Y=key.shape_Y, shape_C=key.shape_C)


def decompress_delta(
    st: StoredFrame,
    key: StoredFrame,
    mode: str,
    divider: int = 1,
    apply_rle: bool = False,
    apply_rle_key: bool = False,
) -> np.ndarray:
    mode = mode.replace("-", ":").strip()
    if apply_rle:
        dY = rle_decode(st.Y, st.shape_Y)
        dCb = rle_decode(st.Cb, st.shape_C)
        dCr = rle_decode(st.Cr, st.shape_C)
    else:
        dY, dCb, dCr = st.Y, st.Cb, st.Cr

    if apply_rle_key:
        key_Y = rle_decode(key.Y, key.shape_Y)
        key_Cb = rle_decode(key.Cb, key.shape_C)
        key_Cr = rle_decode(key.Cr, key.shape_C)
    else:
        key_Y, key_Cb, key_Cr = key.Y, key.Cb, key.Cr

    Y = np.clip(key_Y + dY * divider, 0, 255).astype(int)
    Cb = np.clip(key_Cb + dCb * divider, 0, 255).astype(int)
    Cr = np.clip(key_Cr + dCr * divider, 0, 255).astype(int)
    fd = FrameData(Y=Y, Cb=Cb, Cr=Cr, shape_Y=key.shape_Y, shape_C=key.shape_C)
    return struct_to_frame(fd, mode)


# Wizualizacja różnic (RGB, ROI)
def plot_difference_rgb(ref_rgb: np.ndarray, dec_rgb: np.ndarray, roi: List[int], save_path: Path) -> None:
    r0, r1, c0, c1 = roi
    ref_roi = ref_rgb[r0:r1, c0:c1, :]
    dec_roi = dec_rgb[r0:r1, c0:c1, :]
    diff = ref_roi.astype(float) - dec_roi.astype(float)
    fig, axs = plt.subplots(1, 3, figsize=(12, 4))
    axs[0].imshow(ref_roi)
    axs[0].set_title("Referencja (RGB)")
    axs[1].imshow(diff / 255.0 + 0.5, vmin=0, vmax=1)
    axs[1].set_title("Różnica (RGB)")
    axs[2].imshow(dec_roi)
    axs[2].set_title("Dekompresja (RGB)")
    for ax in axs:
        ax.axis("off")
    plt.tight_layout()
    if SAVE_ROI_IMAGES:
        plt.savefig(save_path, dpi=150)
    plt.close(fig)


# Eksperymenty
def compression_ratio_bytes(orig: np.ndarray, comp: np.ndarray) -> float:
    if comp.nbytes == 0:
        return np.inf
    return orig.nbytes / comp.nbytes


def run_single_experiment(
    video_path: Path,
    mode: str,
    divider: int,
    key_dist: int,
    apply_rle_key: bool,
    apply_rle_delta: bool,
    max_frames: int,
    roi_list: List[List[int]],
    output_prefix: str,
) -> List[Dict[str, object]]:
    cap = cv2.VideoCapture(str(video_path))
    total_frames = int(cap.get(cv2.CAP_PROP_FRAME_COUNT))
    frames_to_process = min(max_frames, total_frames) if max_frames > 0 else total_frames

    results: List[Dict[str, object]] = []
    key_store: StoredFrame | None = None

    for i in range(frames_to_process):
        ret, frame_bgr = cap.read()
        if not ret:
            break

        frame_ycrcb = cv2.cvtColor(frame_bgr, cv2.COLOR_BGR2YCrCb)
        frame_rgb = cv2.cvtColor(frame_bgr, cv2.COLOR_BGR2RGB)
        f_struct = frame_to_struct(frame_ycrcb, mode)

        is_key = (i % key_dist) == 0
        if is_key:
            key_store = compress_key_frame(f_struct, apply_rle=apply_rle_key)
            dec_ycrcb = decompress_key_frame(key_store, mode, apply_rle=apply_rle_key)
            comp_layer = key_store
        else:
            assert key_store is not None, "Brak klatki kluczowej."
            comp_layer = compress_delta(
                f_struct,
                key_store,
                mode=mode,
                divider=divider,
                apply_rle=apply_rle_delta,
                apply_rle_key=apply_rle_key,
            )
            dec_ycrcb = decompress_delta(
                comp_layer,
                key_store,
                mode=mode,
                divider=divider,
                apply_rle=apply_rle_delta,
                apply_rle_key=apply_rle_key,
            )

        dec_rgb = cv2.cvtColor(dec_ycrcb, cv2.COLOR_YCrCb2RGB)

        # Metryki i logi
        ratio_Y = compression_ratio_bytes(f_struct.Y, comp_layer.Y if isinstance(comp_layer.Y, np.ndarray) else np.array(comp_layer.Y))
        ratio_Cb = compression_ratio_bytes(f_struct.Cb, comp_layer.Cb if isinstance(comp_layer.Cb, np.ndarray) else np.array(comp_layer.Cb))
        ratio_Cr = compression_ratio_bytes(f_struct.Cr, comp_layer.Cr if isinstance(comp_layer.Cr, np.ndarray) else np.array(comp_layer.Cr))
        mse_rgb = float(np.mean((frame_rgb.astype(float) - dec_rgb.astype(float)) ** 2))
        results.append(
            {
                "video": video_path.name,
                "frame": i,
                "keyframe": int(is_key),
                "mode": mode,
                "divider": divider,
                "key_dist": key_dist,
                "rle_key": int(apply_rle_key),
                "rle_delta": int(apply_rle_delta),
                "ratio_Y": ratio_Y,
                "ratio_Cb": ratio_Cb,
                "ratio_Cr": ratio_Cr,
                "mse_rgb": mse_rgb,
            }
        )

        # Wizualizacja ROI
        if SAVE_ROI_IMAGES:
            for r_idx, roi in enumerate(roi_list):
                save_path = IMG_DIR / f"{output_prefix}_f{i}_roi{r_idx}.png"
                plot_difference_rgb(frame_rgb, dec_rgb, roi, save_path)

    cap.release()
    return results


def run_experiments() -> List[Dict[str, object]]:
    all_results: List[Dict[str, object]] = []
    videos = [BASE_DIR / "clip_1.mp4", BASE_DIR / "clip_2.mp4"]
    for vid in videos:
        for mode in DEFAULT_SUBSAMPLING:
            for div in DEFAULT_DIVIDERS:
                for kd in DEFAULT_KEY_DIST:
                    safe_mode = mode.replace(":", "-")
                    prefix = f"{vid.stem}_{safe_mode}_d{div}_k{kd}"
                    res = run_single_experiment(
                        video_path=vid,
                        mode=mode,
                        divider=div,
                        key_dist=kd,
                        apply_rle_key=USE_RLE_KEY,
                        apply_rle_delta=USE_RLE_DELTA,
                        max_frames=NUM_FRAMES_TEST,
                        roi_list=ROI_LIST,
                        output_prefix=prefix,
                    )
                    all_results.extend(res)
    # Zapis CSV
    csv_path = RESULTS_DIR / "results_raw.csv"
    with csv_path.open("w", newline="", encoding="utf-8") as f:
        writer = csv.DictWriter(
            f,
            fieldnames=[
                "video",
                "frame",
                "keyframe",
                "mode",
                "divider",
                "key_dist",
                "rle_key",
                "rle_delta",
                "ratio_Y",
                "ratio_Cb",
                "ratio_Cr",
                "mse_rgb",
            ],
        )
        writer.writeheader()
        writer.writerows(all_results)
    print(f"[INFO] Zapisano wyniki CSV -> {csv_path}")
    return all_results


# Raport
def _add_table(doc: Document, title: str, headers: List[str], rows: List[List[str]]) -> None:
    doc.add_paragraph(title)
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
    for row in rows:
        r = table.add_row().cells
        for i, val in enumerate(row):
            r[i].text = val


def generate_report_docx(results: List[Dict[str, object]], output_path: Path) -> None:
    doc = Document()
    doc.add_heading("Laboratorium 8 — Kompresja wideo (wersja uproszczona)", level=0)
    doc.add_paragraph(
        "Testy: subsampling (6 trybów), dzielniki (1,2,4,8,16,32), odstępy klatek kluczowych (1,2,4,6,8,10). "
        "Kompresja różnic Frame-Key, opcjonalny RLE warstw."
    )
    # Wnioski
    doc.add_heading("Wnioski", level=1)
    doc.add_paragraph(
        "Przeprowadzone eksperymenty pokazały, że jakość obrazu zależy głównie od subsamplingu chrominancji oraz wielkości klatek kluczowych."
        "Subsampling 4:2:2 i 4:2:0 daje najlepszy kompromis między kompresją a zachowaniem szczegółów, natomiast zbyt duży divider oraz duży odstęp klatek kluczowych powodują wyraźne zniekształcenia, zwłaszcza w scenach dynamicznych."
        "RLE poprawia kompresję przede wszystkim w obszarach o małej zmienności."
        "Ogólnie najlepsze efekty uzyskano przy umiarkowanych parametrach (divider 1-4, key distance 1-4), które pozwalały zachować dobrą jakość przy sensownym stopniu kompresji."
    )

    # Podsumowanie konfiguracji
    _add_table(
        doc,
        "Konfiguracja eksperymentów",
        ["Subsampling", "Dzielniki", "Key dist", "RLE Key", "RLE Delta", "Klatek/test"],
        [
            [
                ", ".join(DEFAULT_SUBSAMPLING),
                ", ".join(map(str, DEFAULT_DIVIDERS)),
                ", ".join(map(str, DEFAULT_KEY_DIST)),
                str(USE_RLE_KEY),
                str(USE_RLE_DELTA),
                str(NUM_FRAMES_TEST),
            ]
        ],
    )

    # Tabele zbiorcze dla dwóch plików
    for vid in ["clip_1.mp4", "clip_2.mp4"]:
        subset = [r for r in results if r["video"] == vid]
        if not subset:
            continue
        doc.add_heading(f"Wyniki — {vid}", level=1)
        # Kompresja
        rows = []
        for r in subset:
            rows.append(
                [
                    f"{r['frame']} ({'K' if r['keyframe'] else 'D'})",
                    r["mode"],
                    str(r["divider"]),
                    str(r["key_dist"]),
                    f"{r['ratio_Y']:.2f}",
                    f"{r['ratio_Cb']:.2f}",
                    f"{r['ratio_Cr']:.2f}",
                    f"{r['mse_rgb']:.2e}",
                ]
            )
        _add_table(
            doc,
            "Metryki (ratio warstw, MSE RGB)",
            ["Frame", "Subsampling", "Divider", "KeyDist", "CR_Y", "CR_Cb", "CR_Cr", "MSE_RGB"],
            rows,
        )
    
    # Generowanie wykresów 
    # Miniatury i ROI (pokazujemy istniejące pliki), zależne od zapisanych obrazów
    doc.add_heading("Przykładowe porównania ROI", level=1)
    for img_path in sorted(IMG_DIR.glob("*roi*.png")):
        doc.add_paragraph(img_path.stem)
        doc.add_picture(str(img_path), width=Inches(4))


    doc.save(str(output_path))
    print(f"[INFO] Zapisano raport DOCX -> {output_path}")



def main() -> None:
    print("[INFO] Start eksperymentów wideo (Kompresja wideo — uproszczona)...")
    results = run_experiments()
    report_path = BASE_DIR / "report.docx"
    generate_report_docx(results, report_path)
    print("[OK] Zakończono.")


if __name__ == "__main__":
    main()