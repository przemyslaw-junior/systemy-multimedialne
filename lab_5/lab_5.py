from __future__ import annotations

from pathlib import Path
from typing import List, Tuple, Dict

import numpy as np

try:
    from docx import Document
except Exception:
    Document = None


BASE_DIR = Path(__file__).resolve().parent
IMAGES_DIRS = [BASE_DIR / "images", BASE_DIR / "image", BASE_DIR]


def get_size_bytes(x) -> int:
    if isinstance(x, np.ndarray):
        return int(x.nbytes)
    try:
        import sys
        return int(sys.getsizeof(x))
    except Exception:
        return 0


def pack_header(shape: Tuple[int, ...]) -> np.ndarray:
    arr = np.array((len(shape),), dtype=np.int64)
    return np.concatenate([arr, np.array(shape, dtype=np.int64)], axis=0)


def unpack_header(vec: np.ndarray) -> Tuple[Tuple[int, ...], int]:
    vec = np.asarray(vec).astype(np.int64)
    ndim = int(vec[0])
    shape = tuple(int(s) for s in vec[1:1+ndim])
    return shape, 1 + ndim


def to_int_array(x: np.ndarray) -> np.ndarray:
    if not isinstance(x, np.ndarray):
        x = np.asarray(x)
    return x.astype(int, copy=False)

#  RLE (prosty)

def rle_encode(data: np.ndarray) -> np.ndarray:
    x = to_int_array(data)
    flat = x.flatten()
    header = pack_header(x.shape)
    out: List[int] = []
    n = flat.size
    i = 0
    while i < n:
        v = int(flat[i])
        j = i + 1
        while j < n and int(flat[j]) == v:
            j += 1
        run_len = j - i
        out.append(run_len)
        out.append(v)
        i = j
    payload = np.array(out, dtype=np.int64)
    return np.concatenate([header, payload], axis=0)


def rle_decode(vec: np.ndarray) -> np.ndarray:
    vec = to_int_array(vec)
    shape, start = unpack_header(vec)
    payload = vec[start:]
    out: List[int] = []
    i = 0
    m = payload.size
    while i + 1 < m:
        cnt = int(payload[i]); val = int(payload[i+1])
        out.extend([val] * cnt)
        i += 2
    arr = np.array(out, dtype=int)
    return arr.reshape(shape)

#  ByteRun (PackBits-like)

def _run_equal(x: np.ndarray, start: int) -> int:
    n = x.size
    v = x[start]
    j = start + 1
    while j < n and x[j] == v:
        j += 1
    return j - start


def _run_literal(x: np.ndarray, start: int) -> int:
    n = x.size
    if start >= n:
        return 0
    j = start
    while j + 1 < n and x[j] != x[j+1]:
        j += 1
    if j + 1 >= n:
        return n - start
    return (j - start) + 1


def byterun_encode(data: np.ndarray) -> np.ndarray:
    x = to_int_array(data)
    flat = x.flatten()
    header = pack_header(x.shape)
    out: List[int] = []
    n = flat.size
    i = 0
    while i < n:
        eq_len = _run_equal(flat, i)
        if eq_len >= 2:
            v = int(flat[i])
            left = eq_len
            while left > 0:
                chunk = min(left, 128)
                out.append(- (chunk - 1)) 
                out.append(v)
                left -= chunk
            i += eq_len
        else:
            lit_len = _run_literal(flat, i)
            left = lit_len
            j = i
            while left > 0:
                chunk = min(left, 128)
                out.append(chunk - 1)
                for t in range(chunk):
                    out.append(int(flat[j + t]))
                j += chunk
                left -= chunk
            i += lit_len

    payload = np.array(out, dtype=np.int64)
    return np.concatenate([header, payload], axis=0)


def byterun_decode(vec: np.ndarray) -> np.ndarray:
    vec = to_int_array(vec)
    shape, start = unpack_header(vec)
    payload = vec[start:]
    out: List[int] = []
    i = 0
    m = payload.size
    while i < m:
        control = int(payload[i]); i += 1
        if control < 0:
            if i >= m:
                break
            v = int(payload[i]); i += 1
            count = (-control) + 1
            out.extend([v] * count)
        else:
            count = control + 1
            if i + count > m:
                count = max(0, m - i)
            for k in range(count):
                out.append(int(payload[i + k]))
            i += count
    arr = np.array(out, dtype=int)
    return arr.reshape(shape)

#  Testy i metryki

def compression_stats(original: np.ndarray, compressed: np.ndarray) -> Dict[str, float]:
    orig_b = float(get_size_bytes(original))
    comp_b = float(get_size_bytes(compressed))
    cr = orig_b / comp_b if comp_b > 0 else np.inf
    pct = 100.0 * (comp_b / orig_b) if orig_b > 0 else 0.0
    return {"orig_bytes": orig_b, "comp_bytes": comp_b, "cr": cr, "pct": pct}


def verify_roundtrip(name: str, x: np.ndarray) -> Dict[str, Dict[str, float]]:
    x = to_int_array(x)
    results: Dict[str, Dict[str, float]] = {}

    # RLE
    rle_c = rle_encode(x)
    rle_d = rle_decode(rle_c)
    ok_rle = np.array_equal(x, rle_d)
    if not ok_rle:
        print(f"[ERR] RLE roundtrip failed for {name}")
    results["RLE"] = compression_stats(x, rle_c)

    # ByteRun
    br_c = byterun_encode(x)
    br_d = byterun_decode(br_c)
    ok_br = np.array_equal(x, br_d)
    if not ok_br:
        print(f"[ERR] ByteRun roundtrip failed for {name}")
    results["ByteRun"] = compression_stats(x, br_c)

    return results

#  Raport

def build_report(results_arrays: Dict[str, Dict[str, Dict[str, float]]],
                 results_images: Dict[str, Dict[str, Dict[str, float]]]):
    if Document is None:
        print("[INFO] python-docx nie jest dostępne - pomijam raport.")
        return
    doc = Document()
    doc.add_heading("Laboratorium 5 — Kompresja bezstratna (RLE, ByteRun)", level=0)
    doc.add_paragraph(
        "Celem jest implementacja i porównanie kompresji bezstratnej RLE i ByteRun dla tablic NumPy (obrazy/wekory), "
        "z zachowaniem informacji o kształcie wewnątrz jednego wektora danych skompresowanych."
    )

    doc.add_heading("Uwagi dot. pamięci", level=1)
    doc.add_paragraph(
        "Nagłówek skompresowanej postaci zawiera liczbę wymiarów i wymiary oryginalnej tablicy: [ndim, dim0, dim1, ...]. "
        "Dalej następuje strumień danych skompresowanych (RLE: pary [count, value]; ByteRun: naprzemienne znaczniki i wartości zgodnie z opisem)."
    )

    def write_table(title: str, res: Dict[str, Dict[str, float]]):
        doc.add_paragraph(title)
        t = doc.add_table(rows=1, cols=5)
        h = t.rows[0].cells
        h[0].text = "Algorytm"; h[1].text = "rozmiar [B]"; h[2].text = "kompr. [B]"; h[3].text = "CR"; h[4].text = "%"
        for alg, s in res.items():
            r = t.add_row().cells
            r[0].text = alg
            r[1].text = f"{s['orig_bytes']:.0f}"
            r[2].text = f"{s['comp_bytes']:.0f}"
            r[3].text = f"{s['cr']:.3f}"
            r[4].text = f"{s['pct']:.1f}"

    doc.add_heading("Wyniki — zbiory testowe (tablice)", level=1)
    for name, res in results_arrays.items():
        write_table(f"Test: {name}", res)

    doc.add_heading("Wyniki — obrazy (jeśli wczytane)", level=1)
    for name, res in results_images.items():
        write_table(f"Obraz: {name}", res)

    doc.add_heading("Wnioski", level=1)
    doc.add_paragraph(
        "RLE skutecznie kompresuje długie serie identycznych wartości (np. tła, formularze, rysunki), natomiast ByteRun radzi sobie lepiej także z fragmentami bez powtórzeń dzięki trybowi literalnemu. "
        "Efektywność zależy od zawartości: fotografie często kompresują się słabiej. Roundtrip zachowuje identyczność danych."
    )

    path = BASE_DIR / "report.docx"
    doc.save(str(path))
    print("Zapisano raport:", path)

#  I/O obrazów (opcjonalne)

def load_images_from_folder(folder: Path, limit: int = 3) -> Dict[str, np.ndarray]:
    imgs: Dict[str, np.ndarray] = {}
    if not folder.exists():
        print("[INFO] Brak folderu z obrazami:", folder)
        return imgs
    exts = {".png", ".jpg", ".jpeg", ".bmp", ".tif", ".tiff"}
    files = [p for p in sorted(folder.iterdir()) if p.suffix.lower() in exts]
    files = files[:limit]
    if not files:
        print("[INFO] Nie znaleziono obrazów w:", folder)
        return imgs
    reader = None
    try:
        import imageio.v2 as iio  # type: ignore
        reader = "imageio"
    except Exception:
        try:
            from PIL import Image  # type: ignore
            reader = "pil"
        except Exception:
            print("[WARN] Brak imageio/Pillow — pomijam wczytywanie obrazów")
            return imgs

    for p in files:
        try:
            if reader == "imageio":
                arr = to_int_array(iio.imread(str(p)))
            else:
                from PIL import Image
                arr = to_int_array(np.array(Image.open(str(p))))
            if arr.ndim >= 2:
                h, w = arr.shape[0], arr.shape[1]
                if h < 600 or w < 800:
                    print(f"[WARN] {p.name}: rozdzielczosc {w}x{h} mniejsza niz 800x600")
            imgs[p.name] = arr
            print(f"[IMG] Wczytano {p.name} shape={arr.shape}")
        except Exception as e:
            print(f"[ERR] Nie można wczytać {p.name}: {e}")
    return imgs


def load_images_from_candidates(folders: list[Path], limit: int = 3) -> Dict[str, np.ndarray]:
    merged: Dict[str, np.ndarray] = {}
    for f in folders:
        if len(merged) >= limit:
            break
        imgs = load_images_from_folder(f, limit=max(0, limit - len(merged)))
        for name, arr in imgs.items():
            if name in merged:
                continue
            merged[name] = arr
            if len(merged) >= limit:
                break
    if not merged:
        print("[INFO] Nie znaleziono obrazow w images/image/biezacym katalogu lab_5")
    return merged


def main():
    tests: Dict[str, np.ndarray] = {
        "rep_short": np.array([1,1,1,1,2,1,1,1,1,2,1,1,1,1]),
        "cycle": np.array([1,2,3,1,2,3,1,2,3]),
        "mix": np.array([5,1,5,1,5,5,1,1,5,5,1,1,5]),
        "neg_vals": np.array([-1,-1,-1,-5,-5,-3,-4,-2,1,2,2,1]),
        "zeros520": np.zeros((1,520), dtype=int),
        "range521": np.arange(0,521,1),
        "eye7": np.eye(7, dtype=int),
        "stack_eye": np.dstack([np.eye(7, dtype=int), np.eye(7, dtype=int), np.eye(7, dtype=int)]),
        "ones10": np.ones((1,1,1,1,1,1,10), dtype=int),
    }

    results_arrays: Dict[str, Dict[str, Dict[str, float]]] = {}
    for name, arr in tests.items():
        results_arrays[name] = verify_roundtrip(name, arr)

    images = load_images_from_candidates(IMAGES_DIRS, limit=3)
    results_images: Dict[str, Dict[str, Dict[str, float]]] = {}
    for name, arr in images.items():
        results_images[name] = verify_roundtrip(name, arr)

    def print_stats(title: str, res: Dict[str, Dict[str, float]]):
        print(f"\n== {title} ==")
        for alg, s in res.items():
            print(f"{alg}: orig={s['orig_bytes']:.0f}B, comp={s['comp_bytes']:.0f}B, CR={s['cr']:.3f}, {s['pct']:.1f}%")

    for name, res in results_arrays.items():
        print_stats(f"Test {name}", res)
    for name, res in results_images.items():
        print_stats(f"Obraz {name}", res)

    build_report(results_arrays, results_images)


if __name__ == "__main__":
    main()