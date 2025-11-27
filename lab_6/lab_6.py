from __future__ import annotations

import csv
from pathlib import Path
from typing import Callable, Dict, Iterable, List, Tuple

import matplotlib.pyplot as plt
import numpy as np
from matplotlib.backends.backend_pdf import PdfPages
from matplotlib import rcParams

try:
    from docx import Document
    from docx.shared import Inches
except Exception:
    Document = None
    Inches = None

BASE_DIR = Path(__file__).parent
RESULTS_DIR = BASE_DIR / "results"
PLOTS_DIR = RESULTS_DIR / "plots"
AUDIO_DIR = RESULTS_DIR / "audio"
TABLE_DIR = RESULTS_DIR / "tables"

HEAD_COLOR = "#1f497d"
BODY_FONT = "DejaVu Sans"
HEAD_SIZE = 18
SUBHEAD_SIZE = 12
BODY_SIZE = 10
FIGSIZE_A4 = (8.27, 11.69)  # w calach


# Generic helpers

def ensure_dirs() -> None:
    for d in (RESULTS_DIR, PLOTS_DIR, AUDIO_DIR, TABLE_DIR):
        d.mkdir(parents=True, exist_ok=True)


def quantize_uniform(x: np.ndarray, bits: int = 8, vmin: float = -1.0, vmax: float = 1.0) -> np.ndarray:
    x = np.asarray(x, dtype=np.float64)
    levels = 2 ** bits
    step = (vmax - vmin) / (levels - 1)
    clipped = np.clip(x, vmin, vmax)
    quantized = np.round((clipped - vmin) / step) * step + vmin
    return np.clip(quantized, vmin, vmax).astype(np.float32)


# A-law and μ-law companders (vectorised, logical indexing)

def a_law_encode(x: np.ndarray, A: float = 87.6) -> np.ndarray:
    """Vectorised A-law companding; input is clipped to [-1, 1]."""
    x = np.asarray(x, dtype=np.float64)
    x_clip = np.clip(x, -1.0, 1.0)
    abs_x = np.abs(x_clip)
    sign_x = np.sign(x_clip)
    denom = np.log1p(A)
    small = abs_x < (1.0 / A)

    y = np.empty_like(x_clip)
    y[small] = sign_x[small] * (A * abs_x[small]) / denom
    y[~small] = sign_x[~small] * (np.log1p(A * abs_x[~small])) / denom
    return y.astype(np.float32)


def a_law_decode(y: np.ndarray, A: float = 87.6) -> np.ndarray:
    """Inverse of A-law companding; symmetric around zero."""
    y = np.asarray(y, dtype=np.float64)
    y_clip = np.clip(y, -1.0, 1.0)
    abs_y = np.abs(y_clip)
    sign_y = np.sign(y_clip)
    denom = np.log1p(A)
    threshold = 1.0 / denom

    x = np.empty_like(y_clip)
    small = abs_y < threshold
    x[small] = sign_y[small] * (abs_y[small] * denom) / A
    x[~small] = sign_y[~small] * (np.expm1(abs_y[~small] * denom) / A)
    return np.clip(x, -1.0, 1.0).astype(np.float32)


def mu_law_encode(x: np.ndarray, mu: float = 255.0) -> np.ndarray:
    """Vectorised μ-law companding with logical indexing."""
    x = np.asarray(x, dtype=np.float64)
    x_clip = np.clip(x, -1.0, 1.0)
    abs_x = np.abs(x_clip)
    sign_x = np.sign(x_clip)
    y = sign_x * np.log1p(mu * abs_x) / np.log1p(mu)
    return y.astype(np.float32)


def mu_law_decode(y: np.ndarray, mu: float = 255.0) -> np.ndarray:
    """Inverse of μ-law companding."""
    y = np.asarray(y, dtype=np.float64)
    y_clip = np.clip(y, -1.0, 1.0)
    abs_y = np.abs(y_clip)
    sign_y = np.sign(y_clip)
    x = sign_y * (np.expm1(abs_y * np.log1p(mu)) / mu)
    return np.clip(x, -1.0, 1.0).astype(np.float32)


def a_law_compress(x: np.ndarray, bits: int = 8, A: float = 87.6) -> np.ndarray:
    """A-law companding followed by uniform quantization."""
    return quantize_uniform(a_law_encode(x, A), bits, vmin=-1.0, vmax=1.0)


def a_law_decompress(q: np.ndarray, A: float = 87.6) -> np.ndarray:
    """A-law de-companding without additional quantization."""
    return a_law_decode(q, A)


def mu_law_compress(x: np.ndarray, bits: int = 8, mu: float = 255.0) -> np.ndarray:
    """μ-law companding followed by uniform quantization."""
    return quantize_uniform(mu_law_encode(x, mu), bits, vmin=-1.0, vmax=1.0)


def mu_law_decompress(q: np.ndarray, mu: float = 255.0) -> np.ndarray:
    """μ-law de-companding without additional quantization."""
    return mu_law_decode(q, mu)


# DPCM (Differential Pulse Code Modulation)

def _predict_last(x_prev: np.ndarray, n: int | None = None) -> float:
    """Predictor returning the last reconstructed sample (or 0 for empty history)."""
    return float(x_prev[-1]) if x_prev.size else 0.0


def _predict_mean(x_prev: np.ndarray, n: int = 3) -> float:
    """Predictor returning the mean of the last n reconstructed samples."""
    if x_prev.size == 0:
        return 0.0
    return float(np.mean(x_prev[-n:]))


def dpcm_encode(x: np.ndarray, bits: int = 8) -> np.ndarray:
    """Plain DPCM using previous reconstructed sample as predictor."""
    x = np.asarray(x, dtype=np.float64)
    diffs = np.empty_like(x)
    prev = 0.0
    vmin, vmax = -1.0, 1.0
    step = (vmax - vmin) / (2**bits - 1)
    for i, sample in enumerate(x):
        diff = sample - prev
        diff = vmin if diff < vmin else diff
        diff = vmax if diff > vmax else diff
        qdiff = round((diff - vmin) / step) * step + vmin
        qdiff = vmin if qdiff < vmin else qdiff
        qdiff = vmax if qdiff > vmax else qdiff
        diffs[i] = qdiff
        prev = prev + qdiff
    return diffs.astype(np.float32)


def dpcm_decode(diffs: np.ndarray) -> np.ndarray:
    """Reconstruct signal from DPCM differences (no extra predictor)."""
    diffs = np.asarray(diffs, dtype=np.float64)
    x_hat = np.empty_like(diffs)
    prev = 0.0
    for i, d in enumerate(diffs):
        prev = prev + d
        x_hat[i] = prev
    return np.clip(x_hat, -1.0, 1.0).astype(np.float32)


def dpcm_encode_pred(
    x: np.ndarray,
    bits: int = 8,
    predictor: Callable[[np.ndarray, int], float] = _predict_mean,
    n: int = 3,
) -> np.ndarray:
    x = np.asarray(x, dtype=np.float64)
    diffs = np.empty_like(x)
    recon_local = np.empty_like(x)
    vmin, vmax = -1.0, 1.0
    step = (vmax - vmin) / (2**bits - 1)
    for i, sample in enumerate(x):
        pred = predictor(recon_local[:i], n) if i > 0 else 0.0
        diff = sample - pred
        diff = vmin if diff < vmin else diff
        diff = vmax if diff > vmax else diff
        qdiff = round((diff - vmin) / step) * step + vmin
        qdiff = vmin if qdiff < vmin else qdiff
        qdiff = vmax if qdiff > vmax else qdiff
        diffs[i] = qdiff
        recon_local[i] = pred + qdiff
    return diffs.astype(np.float32)


def dpcm_decode_pred(
    diffs: np.ndarray,
    predictor: Callable[[np.ndarray, int], float] = _predict_mean,
    n: int = 3,
) -> np.ndarray:
    diffs = np.asarray(diffs, dtype=np.float64)
    recon = np.empty_like(diffs)
    for i, d in enumerate(diffs):
        pred = predictor(recon[:i], n) if i > 0 else 0.0
        recon[i] = pred + d
    return np.clip(recon, -1.0, 1.0).astype(np.float32)


# Evaluation utilities

def mse(x: np.ndarray, y: np.ndarray) -> float:
    """Mean squared error."""
    x = np.asarray(x, dtype=np.float64)
    y = np.asarray(y, dtype=np.float64)
    return float(np.mean(np.square(x - y)))


def snr_db(x: np.ndarray, x_hat: np.ndarray, eps: float = 1e-12) -> float:
    """Signal-to-noise ratio in dB."""
    x = np.asarray(x, dtype=np.float64)
    x_hat = np.asarray(x_hat, dtype=np.float64)
    num = np.sum(np.square(x))
    den = np.sum(np.square(x - x_hat)) + eps
    return float(10.0 * np.log10(num / den))


def subjective_grade(snr_value: float) -> str:
    if snr_value >= 35:
        return "transparent"
    if snr_value >= 25:
        return "good"
    if snr_value >= 18:
        return "audible artifacts"
    if snr_value >= 12:
        return "noisy but intelligible"
    return "poor / distorted"


def compander_validation_plots(bits: int = 8, out_dir: Path = PLOTS_DIR) -> Path:
    out_dir.mkdir(parents=True, exist_ok=True)
    x = np.linspace(-1, 1, 4000)
    companders = [
        ("A-law", a_law_encode, a_law_decode),
        ("μ-law", mu_law_encode, mu_law_decode),
    ]

    fig, axes = plt.subplots(len(companders), 3, figsize=(13, 7), sharey=False)
    zooms = [
        ("full range", (-1.05, 1.05)),
        ("extremes", (0.8, 1.0)),
        ("near zero", (-0.05, 0.05)),
    ]

    for row, (name, enc_fn, dec_fn) in enumerate(companders):
        y_enc = enc_fn(x)
        y_dec = dec_fn(y_enc)
        for col, (title, xlim) in enumerate(zooms):
            ax = axes[row, col]
            ax.plot(x, y_enc, label=f"{name} encode", alpha=0.85)
            ax.plot(x, y_dec, label=f"{name} decode", alpha=0.85)
            ax.plot(x, x, "--", color="k", linewidth=1, label="ideal")
            ax.set_xlim(xlim)
            ax.set_ylim(-1.05, 1.05)
            ax.set_title(f"{name} - {title}")
            ax.grid(True, linestyle=":")
            if row == len(companders) - 1:
                ax.set_xlabel("input amplitude")
            if col == 0:
                ax.set_ylabel("output")
            if row == 0 and col == 0:
                ax.legend()

    plt.tight_layout()
    path = out_dir / f"validation_companders_bits{bits}.png"
    plt.savefig(path, dpi=200)
    plt.close(fig)
    return path


def dpcm_validation_plots(bits: int = 6, out_dir: Path = PLOTS_DIR) -> Path:
    out_dir.mkdir(parents=True, exist_ok=True)
    x = np.linspace(0, 1, 1500)
    y = 0.9 * np.sin(np.pi * x * 4)

    y_d = dpcm_decode(dpcm_encode(y, bits))
    y_p = dpcm_decode_pred(
        dpcm_encode_pred(y, bits, predictor=_predict_mean, n=3),
        predictor=_predict_mean,
        n=3,
    )

    fig, axes = plt.subplots(2, 1, figsize=(10, 7))
    axes[0].plot(x, y, label="original")
    axes[0].plot(x, y_d, label="DPCM (pred=prev)")
    axes[0].plot(x, y_p, label="DPCM (pred=mean3)")
    axes[0].set_title("DPCM reconstruction on test sinusoid")
    axes[0].legend()
    axes[0].grid(True, linestyle=":")

    axes[1].plot(x, y - y_d, label="error DPCM")
    axes[1].plot(x, y - y_p, label="error DPCM+pred")
    axes[1].set_title("Reconstruction error")
    axes[1].legend()
    axes[1].grid(True, linestyle=":")
    plt.tight_layout()

    path = out_dir / f"validation_dpcm_bits{bits}.png"
    plt.savefig(path, dpi=200)
    plt.close(fig)
    return path


def waveform_plot(
    x: np.ndarray,
    recon: Dict[str, np.ndarray],
    sr: int,
    title: str,
    out_path: Path,
    nsamples: int = 2000,
) -> None:
    t = np.arange(min(len(x), nsamples)) / float(sr)
    plt.figure(figsize=(10, 6))
    plt.plot(t, x[:nsamples], label="original", linewidth=1.2)
    for name, sig in recon.items():
        plt.plot(t, sig[:nsamples], label=name)
    plt.xlabel("time [s]")
    plt.ylabel("amplitude")
    plt.title(title)
    plt.grid(True, linestyle=":")
    plt.legend()
    plt.tight_layout()
    plt.savefig(out_path, dpi=180)
    plt.close()


def load_audio(path: Path) -> Tuple[int, np.ndarray, float, bool]:
    from scipy.io import wavfile

    sr, data = wavfile.read(path)
    is_int = np.issubdtype(data.dtype, np.integer)
    if is_int:
        norm = float(np.iinfo(data.dtype).max)
        data_f = data.astype(np.float32) / norm
    else:
        norm = 1.0
        data_f = data.astype(np.float32)
    return sr, data_f, norm, is_int


def save_audio(path: Path, sr: int, data: np.ndarray, norm: float, as_int: bool) -> None:
    from scipy.io import wavfile

    if as_int:
        wav_data = np.int16(np.clip(data * norm, -norm, norm - 1))
    else:
        wav_data = data.astype(np.float32)
    wavfile.write(path, sr, wav_data)


def process_method(x: np.ndarray, bits: int, method: str, pred_n: int = 3) -> np.ndarray:
    if method == "a_law":
        return a_law_decompress(a_law_compress(x, bits))
    if method == "mu_law":
        return mu_law_decompress(mu_law_compress(x, bits))
    if method == "dpcm_np1":
        return dpcm_decode_pred(
            dpcm_encode_pred(x, bits, predictor=_predict_mean, n=1),
            predictor=_predict_mean,
            n=1,
        )
    if method == "dpcm_np3":
        return dpcm_decode_pred(
            dpcm_encode_pred(x, bits, predictor=_predict_mean, n=pred_n),
            predictor=_predict_mean,
            n=pred_n,
        )
    raise ValueError(f"Unknown method: {method}")


def run_experiments(
    audio_files: Iterable[Path],
    bits_list: Iterable[int],
    out_audio_dir: Path = AUDIO_DIR,
    out_plot_dir: Path = PLOTS_DIR,
) -> List[Dict[str, object]]:
    ensure_dirs()
    out_audio_dir.mkdir(parents=True, exist_ok=True)
    out_plot_dir.mkdir(parents=True, exist_ok=True)

    methods = ["a_law", "mu_law", "dpcm_np1", "dpcm_np3"]
    results: List[Dict[str, object]] = []

    for audio_path in audio_files:
        sr, data_f, norm, is_int = load_audio(audio_path)
        if data_f.ndim == 1:
            data_f = data_f[:, None]

        for bits in bits_list:
            recon_channels: Dict[str, np.ndarray] = {m: np.zeros_like(data_f) for m in methods}

            for ch in range(data_f.shape[1]):
                xch = data_f[:, ch]
                for m in methods:
                    recon_channels[m][:, ch] = process_method(xch, bits, method=m, pred_n=3)

            for m in methods:
                save_audio(
                    out_audio_dir / f"{audio_path.stem}_{m}_{bits}bit.wav",
                    sr,
                    recon_channels[m].squeeze(),
                    norm,
                    as_int=is_int,
                )

            x_mono = np.mean(data_f, axis=1)
            for m in methods:
                y_mono = np.mean(recon_channels[m], axis=1)
                snr_value = snr_db(x_mono, y_mono)
                mse_value = mse(x_mono, y_mono)
                results.append(
                    {
                        "file": audio_path.name,
                        "bits": bits,
                        "method": m,
                        "snr_db": snr_value,
                        "mse": mse_value,
                        "subjective": subjective_grade(snr_value),
                    }
                )

            wav_plot_path = out_plot_dir / f"waveform_{audio_path.stem}_{bits}bit.png"
            waveform_plot(
                x=np.mean(data_f, axis=1),
                recon={m: np.mean(recon_channels[m], axis=1) for m in methods},
                sr=sr,
                title=f"{audio_path.name} @ {bits} bits",
                out_path=wav_plot_path,
                nsamples=min(5000, data_f.shape[0]),
            )

    return results



def save_results_csv(results: List[Dict[str, object]], path: Path = TABLE_DIR / "results.csv") -> None:
    if not results:
        return
    path.parent.mkdir(parents=True, exist_ok=True)
    fieldnames = ["file", "bits", "method", "snr_db", "mse", "subjective"]
    with path.open("w", newline="", encoding="utf-8") as f:
        writer = csv.DictWriter(f, fieldnames=fieldnames)
        writer.writeheader()
        writer.writerows(results)


def recognizability_table(results: List[Dict[str, object]]) -> Dict[str, Dict[int, str]]:
    table: Dict[str, Dict[int, str]] = {}
    for row in results:
        fname = row["file"]
        bits = int(row["bits"])
        grade = str(row["subjective"])
        if fname not in table:
            table[fname] = {}
        # choose best (highest SNR) grade per bit-depth
        if bits not in table[fname]:
            table[fname][bits] = grade
        else:
            existing = table[fname][bits]
            order = ["poor / distorted", "noisy but intelligible", "audible artifacts", "good", "transparent"]
            if order.index(grade) > order.index(existing):
                table[fname][bits] = grade
    return table


def _table_figure(table: Dict[str, Dict[int, str]], bits_list: List[int]) -> Path:
    fig, ax = plt.subplots(figsize=(10, 0.6 * (len(table) + 2)))
    ax.axis("off")
    rows = list(table.keys())
    cell_text = []
    for fname in rows:
        row_vals = []
        for b in bits_list:
            row_vals.append(table[fname].get(b, "n/a"))
        cell_text.append(row_vals)
    col_labels = [f"{b} bit" for b in bits_list]
    tab = ax.table(
        cellText=cell_text,
        rowLabels=rows,
        colLabels=col_labels,
        loc="center",
        cellLoc="center",
    )
    tab.scale(1.0, 1.5)
    plt.tight_layout()
    img_path = PLOTS_DIR / "recognizability_table.png"
    plt.savefig(img_path, dpi=200)
    plt.close(fig)
    return img_path


def _init_report_style() -> None:
    rcParams.update(
        {
            "font.family": BODY_FONT,
            "font.size": BODY_SIZE,
            "axes.titlesize": SUBHEAD_SIZE,
            "axes.labelsize": BODY_SIZE,
        }
    )


def _text_block(y: float, text: str, size: int = BODY_SIZE, color: str = "black", weight: str = "normal") -> float:
    plt.text(0.05, y, text, ha="left", va="top", wrap=True, fontsize=size, color=color, fontweight=weight)
    return y - 0.05


def _add_cover_page(pdf: PdfPages, title: str, subtitle: str, body: str) -> None:
    fig = plt.figure(figsize=FIGSIZE_A4)
    plt.axis("off")
    y = 0.92
    plt.text(0.05, y, title, ha="left", va="top", fontsize=HEAD_SIZE, color=HEAD_COLOR, fontweight="bold")
    y -= 0.04
    plt.plot([0.05, 0.6], [y, y], color=HEAD_COLOR, linewidth=1.5)
    y -= 0.06
    plt.text(0.05, y, subtitle, ha="left", va="top", fontsize=SUBHEAD_SIZE, color=HEAD_COLOR, fontweight="bold")
    y -= 0.06
    plt.text(0.05, y, body, ha="left", va="top", wrap=True, fontsize=BODY_SIZE, color="black", linespacing=1.4)
    pdf.savefig(fig)
    plt.close(fig)


def _add_figure_page(pdf: PdfPages, img_path: Path, caption: str, fig_number: int) -> None:
    fig = plt.figure(figsize=FIGSIZE_A4)
    plt.axis("off")
    img = plt.imread(img_path)
    plt.imshow(img)
    plt.axis("off")
    plt.figtext(0.05, 0.08, f"Rys. {fig_number}: {caption}", ha="left", va="bottom", fontsize=BODY_SIZE, color="black")
    pdf.savefig(fig)
    plt.close(fig)


def build_report_docx(
    results: List[Dict[str, object]],
    comp_plot: Path,
    dpcm_plot: Path,
    bits_list: List[int],
    path: Path = RESULTS_DIR / "report.docx",
) -> None:
    if Document is None:
        print("[INFO] python-docx niedostępne - pomijam raport DOCX.")
        return
    ensure_dirs()
    doc = Document()
    doc.add_heading("Laboratorium 6 — Kompresja stratna audio (A-law, μ-law, DPCM)", level=0)
    doc.add_paragraph(
        "Celem jest porównanie metod kompresji stratnej audio: A-law, μ-law oraz DPCM z/bez predykcji. "
        "Zbadano wpływ kwantyzacji (8-2 bity) na jakość sygnałów z trzech plików (low/medium/high). "
        "Prezentowane są krzywe kompresji, wyniki SNR/MSE oraz ocena rozpoznawalności sygnału."
    )

    doc.add_heading("Uwagi dot. danych", level=1)
    doc.add_paragraph(
        "Sygnały wczytywane jako float32, kwantyzacja uniform 8..2 bitów. Predykcja w DPCM: średnia z n próbek."
    )

    doc.add_heading("Walidacja kompresji", level=1)
    doc.add_paragraph("Krzywe A-law i μ-law (pełny zakres oraz zbliżenia):")
    if comp_plot.exists() and Inches:
        doc.add_picture(str(comp_plot), width=Inches(6))
    doc.add_paragraph("DPCM - rekonstrukcja i błąd na sinusoidzie testowej:")
    if dpcm_plot.exists() and Inches:
        doc.add_picture(str(dpcm_plot), width=Inches(6))

    doc.add_heading("Tabela rozpoznawalności (8..2 bity)", level=1)
    recog_table = recognizability_table(results)
    table_img = _table_figure(recog_table, bits_list)
    if table_img.exists() and Inches:
        doc.add_picture(str(table_img), width=Inches(5.5))

    doc.add_heading("Wyniki — SNR/MSE", level=1)
    if results:
        t = doc.add_table(rows=1, cols=6)
        h = t.rows[0].cells
        h[0].text = "Plik"
        h[1].text = "Bity"
        h[2].text = "Metoda"
        h[3].text = "SNR [dB]"
        h[4].text = "MSE"
        h[5].text = "Subiektywna"
        for r in results:
            row = t.add_row().cells
            row[0].text = str(r["file"])
            row[1].text = str(r["bits"])
            row[2].text = str(r["method"])
            row[3].text = f"{r['snr_db']:.2f}"
            row[4].text = f"{r['mse']:.2e}"
            row[5].text = str(r.get("subjective", ""))

    doc.add_heading("Wnioski", level=1)
    doc.add_paragraph(
        "A-law/μ-law zachowują wysoką rozpoznawalność przy niższej liczbie bitów. "
        "Predykcja w DPCM poprawia SNR dla wolnozmiennych sygnałów. "
        "Poniżej 4 bitów jakość szybko degraduje się - odsłuch wskazany."
    )

    doc.save(str(path))
    print(f"[OK] Zapisano raport DOCX: {path}")


def run_all() -> None:
    ensure_dirs()
    bits_list = [8, 7, 6, 5, 4, 3, 2]

    default_files = [
        BASE_DIR.parent / "lab_4" / "SING" / "sing_low1.wav",
        BASE_DIR.parent / "lab_4" / "SING" / "sing_medium1.wav",
        BASE_DIR.parent / "lab_4" / "SING" / "sing_high1.wav",
    ]
    audio_files = [p for p in default_files if p.exists()]
    if not audio_files:
        raise FileNotFoundError("No default audio files found; adjust `default_files` in run_all().")

    comp_plot = compander_validation_plots(bits=8, out_dir=PLOTS_DIR)
    dpcm_plot = dpcm_validation_plots(bits=6, out_dir=PLOTS_DIR)
    results = run_experiments(audio_files=audio_files, bits_list=bits_list)
    save_results_csv(results, path=TABLE_DIR / "results.csv")
    # Raport w stylu lab 5 (DOCX)
    build_report_docx(results, comp_plot=comp_plot, dpcm_plot=dpcm_plot, bits_list=bits_list)


if __name__ == "__main__":
    run_all()